"""
產出麥肯錫格式向上報告 Word 檔
輸出路徑：99_工作暫存/BD自動化工作流_向上報告.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

# ── 顏色 ────────────────────────────────────────────────────────────────────
NAVY      = RGBColor(0,  48, 130)      # 麥肯錫深藍
MID_BLUE  = RGBColor(0,  70, 160)
DARK_GREY = RGBColor(64, 64,  64)
BLACK     = RGBColor(0,   0,   0)
WHITE     = RGBColor(255,255, 255)
FILL_BLUE = "EBF3FB"   # 標注框淡藍底
FILL_GREY = "F4F4F4"   # 流程框淡灰底
FILL_NAV  = "003082"   # 表格標題列深藍
FILL_ROW  = "F0F4FA"   # 表格偶數列淡藍

FONT = "Calibri"
OUT  = Path(r"c:\Users\C0378\Desktop\旺來PPT製作\嘟嘟房_提案專案\99_工作暫存\BD自動化工作流_向上報告.docx")

# ── 工具函式 ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, fill_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tc_pr = cell._tc.get_or_add_tcPr()
    # 移除舊的 shd
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    tc_pr.append(shd)

def set_para_border_bottom(para, color_hex="003082", sz="8"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    sz)
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)

def set_cell_padding(cell, top=60, bottom=60, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def no_border_table(table):
    """移除表格框線，僅保留內部整潔感"""
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ["top","left","bottom","right","insideH","insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBdr.append(el)
    tblPr.append(tblBdr)

def thin_border_table(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement("w:tblBorders")
    for side in ["top","left","bottom","right","insideH","insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CCCCCC")
        tblBdr.append(el)
    tblPr.append(tblBdr)

# ── 文字加入輔助 ─────────────────────────────────────────────────────────────
def p_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name  = FONT
    r.font.size  = Pt(20)
    r.font.bold  = True
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(4)
    return p

def p_meta(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}　")
    r1.font.name  = FONT
    r1.font.size  = Pt(10.5)
    r1.font.bold  = True
    r1.font.color.rgb = DARK_GREY
    r2 = p.add_run(value)
    r2.font.name  = FONT
    r2.font.size  = Pt(10.5)
    r2.font.color.rgb = DARK_GREY
    p.paragraph_format.space_after = Pt(2)
    return p

def p_section(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name  = FONT
    r.font.size  = Pt(14)
    r.font.bold  = True
    r.font.color.rgb = NAVY
    set_para_border_bottom(p)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    return p

def p_sub(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name  = FONT
    r.font.size  = Pt(11.5)
    r.font.bold  = True
    r.font.color.rgb = MID_BLUE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def p_body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name   = FONT
    r.font.size   = Pt(11)
    r.font.bold   = bold
    r.font.italic = italic
    r.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(5)
    return p

def p_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(0.6 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after   = Pt(3)
    dot = p.add_run("• ")
    dot.font.name  = FONT
    dot.font.size  = Pt(11)
    dot.font.color.rgb = NAVY
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.font.name  = FONT
        rb.font.size  = Pt(11)
        rb.font.bold  = True
        rb.font.color.rgb = BLACK
    r = p.add_run(text)
    r.font.name  = FONT
    r.font.size  = Pt(11)
    r.font.color.rgb = BLACK
    return p

def p_callout(doc, text, fill=FILL_BLUE):
    tbl = doc.add_table(rows=1, cols=1)
    no_border_table(tbl)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, fill)
    set_cell_padding(cell, 100, 100, 180, 120)
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name   = FONT
    r.font.size   = Pt(11)
    r.font.bold   = True
    r.font.italic = True
    r.font.color.rgb = NAVY
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def p_process(doc, lines):
    tbl  = doc.add_table(rows=1, cols=1)
    no_border_table(tbl)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, FILL_GREY)
    set_cell_padding(cell, 120, 120, 200, 120)
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        r = p.add_run(line)
        r.font.name  = "Courier New"
        r.font.size  = Pt(9.5)
        r.font.color.rgb = DARK_GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def p_spacer(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)
    return p

def make_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    thin_border_table(tbl)
    # 標題列
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, FILL_NAV)
        set_cell_padding(cell, 80, 80, 120, 120)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name  = FONT
        r.font.size  = Pt(10)
        r.font.bold  = True
        r.font.color.rgb = WHITE
    # 資料列
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        fill = FILL_ROW if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, fill)
            set_cell_padding(cell, 60, 60, 120, 120)
            p = cell.paragraphs[0]
            # 加粗第一欄
            bold = (ci == 0)
            r = p.add_run(val)
            r.font.name  = FONT
            r.font.size  = Pt(10)
            r.font.bold  = bold
            r.font.color.rgb = BLACK
    # 欄寬
    if col_widths:
        for i, row in enumerate(tbl.rows):
            for j, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[j])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

# ── 建立文件 ─────────────────────────────────────────────────────────────────
doc = Document()

# 版面設定
for sec in doc.sections:
    sec.top_margin    = Cm(2.54)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.8)
    sec.right_margin  = Cm(2.8)

# 預設字型
doc.styles["Normal"].font.name = FONT
doc.styles["Normal"].font.size = Pt(11)

# ════════════════════════════════════════════════════════════════════════════
# 封面資訊
# ════════════════════════════════════════════════════════════════════════════
p_spacer(doc, 20)
p_title(doc, "開發部 AI 陌生開發工作流")
p_title(doc, "建置成果報告")
p_spacer(doc, 10)

# 分隔線
div = doc.add_paragraph()
div_r = div.add_run("─" * 52)
div_r.font.color.rgb = NAVY
div_r.font.size = Pt(10)
div.paragraph_format.space_after = Pt(10)

p_meta(doc, "報告對象：", "主管")
p_meta(doc, "撰寫部門：", "開發部")
p_meta(doc, "報告日期：", "2026 年 6 月 2 日")
p_meta(doc, "報告性質：", "工作流建置成果說明 ／ 未來效率提升機制報告")
p_spacer(doc, 14)

# ════════════════════════════════════════════════════════════════════════════
# 核心結論（Executive Summary — 麥肯錫金字塔：結論先說）
# ════════════════════════════════════════════════════════════════════════════
p_callout(doc,
    "核心結論：開發部已建立一套 AI 陌生開發自動化工作流，"
    "可將新品牌從「完全陌生」到「提案材料完整就緒」的時間，"
    "從過去約 12 至 15 個工作天縮短至 2.5 至 4 個工作天（節省約 75–80%）。"
    "本套工作流可直接複製至下一個陌生開發目標，無需重新設計。")

# ════════════════════════════════════════════════════════════════════════════
# 一、報告摘要
# ════════════════════════════════════════════════════════════════════════════
p_section(doc, "一、這份報告在說什麼")
p_body(doc,
    "開發部在本次嘟嘟房 BD 專案推進期間，同步完成了一套「陌生開發自動化工作流」的建置。"
    "這套工作流的意義不只是嘟嘟房這一個案子，而是：")
p_callout(doc,
    "未來開發部每次面對新的陌生開發目標（停車場業者、通路品牌、場域合作夥伴），"
    "都可以用這套方法，在幾天內完成從「不認識對方」到「把提案材料送到對方桌上」的全流程。")
p_body(doc, "本報告說明三件事：我們做了什麼、解決了什麼問題、以後能省多少時間。")

# ════════════════════════════════════════════════════════════════════════════
# 二、過去的問題
# ════════════════════════════════════════════════════════════════════════════
p_section(doc, "二、我們以前怎麼做陌生開發？問題在哪？")
p_body(doc, "在沒有這套工作流之前，開發部的陌生開發作業流程如下：")

p_process(doc, [
    "  同仁拿到目標品牌名稱",
    "      ↓",
    "  自己上網查資料（沒有結構，查到哪算哪）",
    "      ↓",
    "  用感覺寫一份 Email 或一頁提案",
    "      ↓",
    "  內容品質高度依賴個人經驗",
    "      ↓",
    "  主管 review 後可能要大幅修改",
    "      ↓",
    "  修改完才能寄出",
])

p_sub(doc, "這個做法有四個根本性問題")
make_table(doc,
    ["問題", "說明"],
    [
        ["調查沒有框架",   "沒有標準的調查項目清單，不同人查出來的資料深度差很多"],
        ["定位判斷靠感覺", "未強制回答「為什麼找這個品牌」「對方能得到什麼」就直接開始寫提案"],
        ["材料品質不穩定", "個人能力越強，材料越好；換人做，品質就不一樣"],
        ["每次從零開始",   "做完一個品牌的開發，下一個品牌又要重頭來過，沒有任何沉澱"],
    ],
    col_widths=[4.5, 11.5]
)

# ════════════════════════════════════════════════════════════════════════════
# 三、我們建立了什麼
# ════════════════════════════════════════════════════════════════════════════
p_section(doc, "三、我們現在建立了什麼？")
p_body(doc,
    "我們把「從陌生到提案」的全流程，拆成 11 個明確的工作節點。"
    "每個節點都有一份「AI 執行指令」，讓 AI 像一個訓練有素的開發助理，"
    "按照我們設計好的邏輯，逐步產出每個環節的文件。")

p_sub(doc, "第一層｜先了解對方（不動筆，先做功課）")
make_table(doc,
    ["節點", "任務名稱", "AI 執行內容", "產出文件"],
    [
        ["節點 −1",   "公開情報調查（DD）",     "針對目標品牌執行 13 個調查區塊，涵蓋公司規模、商業模式、市場地位、決策鏈、合作風險",
                                              "DD 報告 / BD 策略稿 / 場站清單與風險登錄 / 總部接觸策略"],
        ["節點 −0.5", "候選場站資料庫建立",     "針對桃園 / 台中的候選場站建立評分資料庫，根據合作價值、進入難度、供電條件逐站評分",
                                              "Excel 案場資料庫（含 TOP10 推薦與供電分析）"],
    ],
    col_widths=[2.2, 3.5, 6.5, 4.8]
)

p_sub(doc, "第二層｜決定怎麼說（定位先於文案）")
make_table(doc,
    ["節點", "任務名稱", "AI 執行內容", "產出文件"],
    [
        ["節點 0", "資料盤點",     "確認所有文件是否齊全，才能繼續",                                     "盤點確認清單"],
        ["節點 1", "專案定位分析", "強制回答：為什麼找這個品牌？對方能得到什麼？風險由誰承擔？",           "定位分析稿（所有後續文案的地基）"],
    ],
    col_widths=[2.2, 3.5, 6.5, 4.8]
)

p_sub(doc, "第三層｜產出第一份對外材料（OnePager DM）")
make_table(doc,
    ["節點", "任務名稱", "AI 執行內容", "產出文件"],
    [
        ["節點 2", "OnePager 內容稿",  "在一頁以內，回答對方最想知道的 4 個問題（外部商務語氣）", "OnePager.md"],
        ["節點 3", "OnePager DM 設計", "將內容稿轉成可作附件的 A4 設計版面（旺來品牌風格）",      "OnePager_DM.html"],
        ["節點 4", "PNG / PDF 輸出",   "用程式自動輸出設計品質的圖片與文件，可直接作 Email 附件", "PNG 圖片 + PDF 檔案"],
    ],
    col_widths=[2.2, 3.5, 6.5, 4.8]
)

p_sub(doc, "第四層｜接觸對方（從寄信到取得會議）")
make_table(doc,
    ["節點", "任務名稱", "AI 執行內容", "產出文件"],
    [
        ["節點 5", "Cold Mail 工作流",   "產出 3 種主旨選項、正式版與精簡版信件、可貼上 Gmail 的純文字版，附寄送前檢查清單",
                                         "Cold Mail 工作流 / 三版信件 / 追蹤時程表"],
        ["節點 6", "追蹤話術與取得會議", "D+4 追蹤信 / D+10 換角度追蹤信 / D+14 電話逐字稿 / 30 分鐘會議完整腳本 / 常見問題應答表",
                                         "追蹤信件 × 2 + 電話腳本 + 會議腳本"],
    ],
    col_widths=[2.2, 3.5, 6.5, 4.8]
)

p_sub(doc, "第五層｜進入會議後的提案材料")
make_table(doc,
    ["節點", "任務名稱", "AI 執行內容", "產出文件"],
    [
        ["節點 7", "簡版 PPT 大綱", "10 頁，初步對焦會議用，每頁一個核心結論，最後只問對方一件事",             "ShortPPT_Outline.md"],
        ["節點 8", "正式 PPT 大綱", "16 頁，L2 總經理層級，數字驅動，含試點 KPI 與時程表，附完整附錄框架",   "FormalPPT_Outline.md"],
    ],
    col_widths=[2.2, 3.5, 6.5, 4.8]
)

# ════════════════════════════════════════════════════════════════════════════
# 四、嘟嘟房案例
# ════════════════════════════════════════════════════════════════════════════
p_section(doc, "四、實際案例：嘟嘟房專案產出了多少文件？")
p_body(doc, "以本次嘟嘟房案作為完整示範，AI 在這個案子中共產出 22 份文件：")

make_table(doc,
    ["文件類型", "具體文件內容"],
    [
        ["情報調查（5 份）",   "DD 執行指南 / DD 調查報告 / BD 策略稿 / 場站與決策鏈 / 總部接觸策略"],
        ["場站資料庫（1 份）", "Excel 9 個工作表（台中 15 站、桃園 10 站評分、TOP10、供電分析）"],
        ["定位分析（1 份）",   "定位分析稿（為什麼找嘟嘟房 / 四個必答問題 / 試點假設 / 禁止話題清單）"],
        ["OnePager（4 份）",   "內容稿 / 設計版 HTML / 設計版 PNG（2× 高解析） / 設計版 PDF（A4 單頁）"],
        ["Cold Mail（5 份）",  "工作流文件 / 三版主旨 / 正式信 / 精簡信 / 可貼 Email 的純文字版"],
        ["追蹤話術（4 份）",   "D+4 追蹤信 / D+10 追蹤信 / D+14 電話腳本 / 30 分鐘會議完整逐字腳本"],
        ["PPT 大綱（2 份）",   "10 頁簡版大綱 / 16 頁正式提案大綱"],
    ],
    col_widths=[4.0, 13.0]
)

# ════════════════════════════════════════════════════════════════════════════
# 五、AI 的角色
# ════════════════════════════════════════════════════════════════════════════
p_section(doc, "五、AI 在整個過程中扮演什麼角色？")
p_body(doc,
    "很多人聽到「AI」，第一個想法是「幫我生文章」。"
    "這個工作流裡，AI 扮演三個截然不同的角色：")

p_sub(doc, "角色一：資料調查員")
p_body(doc,
    "AI 能在幾十分鐘內完成原本需要 2 至 3 天的公開情報調查，"
    "結果結構化、可存檔、可複查：")
for item in [
    "公司規模與商業模式",
    "市場競爭地位（誰是第一、誰是第二、誰在追上來）",
    "組織決策鏈（誰有決策權、要找哪個層級的人）",
    "風險評估（IPO 敏感期、對方偏好的合作形式）",
    "合作適配性評分（A / B / C 級）",
]:
    p_bullet(doc, item)

p_sub(doc, "角色二：BD 策略師")
p_body(doc,
    "有了情報之後，AI 依照開發部設計的「定位判斷框架」，"
    "強制回答以下五個問題，才能進入文案撰寫，順序不能顛倒：")
p_process(doc, [
    "  1. 為什麼現在要找這個品牌？（而不是六個月後）",
    "  2. 這個品牌為什麼有動機合作？（具體原因，不是模糊的「市場大」）",
    "  3. 對方能得到什麼具體的東西？（場站加值 / 品牌升級 / 數據 / 上市故事）",
    "  4. 如果對方說「這很麻煩」，我們怎麼回答？",
    "  5. 什麼話題絕對不能在第一次接觸時提？",
])

p_sub(doc, "角色三：文件生產員")
p_body(doc,
    "定位確定後，AI 根據統一的邏輯框架，依序生產每一份文件，"
    "且每份文件之間的邏輯是互相連貫的，不會「Cold Mail 說一套，PPT 說另一套」：")
p_process(doc, [
    "  定位分析  →  為 OnePager 定調",
    "  OnePager  →  為 Cold Mail 提供附件與核心主張",
    "  Cold Mail →  為追蹤信與電話腳本建立延伸語境",
    "  追蹤話術  →  為 30 分鐘會議腳本建立接話基礎",
    "  會議腳本  →  為 PPT 設計決定頁面邏輯順序",
])

p_sub(doc, "哪些事情 AI 不能做？（同樣重要）")
make_table(doc,
    ["AI 能做的", "仍然需要人的判斷"],
    [
        ["公開資料的結構化調查",    "確認 AI 調查結果是否符合現況（部分公開資料可能已過期）"],
        ["框架內的文件依序生產",    "每個節點完成後的「要不要過」判斷"],
        ["按照邏輯產出文案初稿",    "最後寄信前，確認主旨與語氣是否適合這個具體對象"],
        ["設計版面的程式技術輸出",  "場站條件的現勘，仍須實地確認"],
    ],
    col_widths=[5.5, 11.5]
)
p_callout(doc,
    "設計原則：每個節點完成後，AI 必須停下來等待人員確認，不能自動執行下一步。"
    "人的判斷沒有被跳過，只是不再被繁瑣的文件生產工作佔用時間。",
    fill="FFF3E0")

# ════════════════════════════════════════════════════════════════════════════
# 六、效率估算
# ════════════════════════════════════════════════════════════════════════════
p_section(doc, "六、效率提升的具體估算")
p_sub(doc, "傳統做法 vs. 新工作流（時間比較）")
make_table(doc,
    ["工作項目", "傳統做法（人工）", "新工作流（AI 輔助）", "節省幅度"],
    [
        ["陌生品牌公開情報調查",            "2 至 3 天",  "2 至 4 小時", "約 80%"],
        ["候選場站評分資料庫建立",          "1 至 2 天",  "3 至 6 小時", "約 70%"],
        ["定位分析稿（為什麼找這個品牌）",  "半天",       "1 至 2 小時", "約 70%"],
        ["OnePager 內容稿 + 設計版",        "1 至 3 天",  "半天",        "約 80%"],
        ["Cold Mail 三版 + 主旨選項",       "1 天",       "2 至 3 小時", "約 75%"],
        ["追蹤信 + 電話腳本 + 會議腳本",    "2 天",       "半天",        "約 80%"],
        ["簡版 PPT 大綱（10 頁）",          "1 天",       "1 至 2 小時", "約 80%"],
        ["正式 PPT 大綱（16 頁）",          "2 天",       "半天",        "約 75%"],
        ["合計：一個品牌從零到就緒",         "12 至 15 天","2.5 至 4 天", "約 75–80%"],
    ],
    col_widths=[5.5, 3.5, 4.0, 3.0]
)
p_body(doc, "備注：以上為「第一次使用工作流、邊學習邊建置」的時間估算。流程熟悉後，預計可進一步壓縮至 1.5 至 2 個工作天。",
       italic=True)

p_sub(doc, "品質提升（時間之外的另一個維度）")
for bp, rest in [
    ("不再靠感覺寫定位：", "每個品牌都必須先回答四個強制性問題，才能進入文案"),
    ("文件邏輯互相連貫：", "不會出現「Cold Mail 說一套，PPT 說另一套」的狀況"),
    ("新人也能執行：",     "工作流節點明確，照著做就能產出接近資深業務水準的材料"),
    ("每次開發都有沉澱：", "每個品牌的文件都留存在對應資料夾，日後可回頭審視當初的策略判斷"),
]:
    p_bullet(doc, rest, bold_prefix=bp)

# ════════════════════════════════════════════════════════════════════════════
# 七、延伸應用
# ════════════════════════════════════════════════════════════════════════════
p_section(doc, "七、未來可以直接套用的範疇")
p_body(doc,
    "這套工作流目前以停車場業者為主要情境設計，"
    "但核心邏輯可延伸至其他陌生開發場景：")
make_table(doc,
    ["開發場景", "可複用程度", "備注"],
    [
        ["其他停車場品牌（台灣聯通、快速停車等）", "直接複製", "修改品牌名稱即可使用"],
        ["便利商店 / 超市等通路品牌",              "大部分可用", "調整 DD 調查的產業邏輯區塊"],
        ["企業園區 / 商辦大樓",                    "大部分可用", "場站型態與停車場類似"],
        ["社區管委會 / 住宅開發商",                "部分可用", "需調整決策鏈判斷邏輯"],
    ],
    col_widths=[5.5, 3.5, 8.0]
)

# ════════════════════════════════════════════════════════════════════════════
# 八、本質
# ════════════════════════════════════════════════════════════════════════════
p_section(doc, "八、這套工作流的本質是什麼？")
p_callout(doc,
    "我們把「一個有經驗的 BD 主管腦子裡的判斷邏輯」，"
    "整理成可以重複執行的步驟，然後讓 AI 按照這個步驟工作。")
p_body(doc, "這不是「讓 AI 取代業務」。")
p_body(doc, "這是「讓業務人員的時間花在判斷與接觸，而不是重複性的文件生產」。")
p_spacer(doc, 4)
p_body(doc,
    "每一份最終寄出的信、最終帶去開會的 PPT，仍然需要業務人員的判斷與確認。"
    "這套工作流做的事情，是把「準備材料」這件事的時間，從 12 天壓到 3 天以內。")

# ════════════════════════════════════════════════════════════════════════════
# 九、建議後續行動
# ════════════════════════════════════════════════════════════════════════════
p_section(doc, "九、建議後續行動")
make_table(doc,
    ["優先級", "建議事項", "預期效益"],
    [
        ["P0 — 立即",   "確認嘟嘟房 Cold Mail 發送時間，啟動第一波接觸",
                        "驗證工作流產出的材料在真實場景中的效果"],
        ["P1 — 本月",   "選定下一個目標品牌，用這套工作流跑第二輪",
                        "預計可在首輪基礎上進一步提速至 1.5 至 2 天"],
        ["P2 — 本季",   "考慮將 BD 自動化工作流納入開發部新人培訓材料",
                        "降低新人學習曲線，縮短獨立作業所需時間"],
        ["P3 — 每季",   "定期更新各品牌案場資料庫（場站地址、管理層異動、市場地位）",
                        "確保資料庫與市場現況同步，避免使用過期資訊"],
    ],
    col_widths=[2.8, 7.2, 7.0]
)

# ════════════════════════════════════════════════════════════════════════════
# 頁尾
# ════════════════════════════════════════════════════════════════════════════
p_spacer(doc, 16)
div2 = doc.add_paragraph()
div2_r = div2.add_run("─" * 52)
div2_r.font.color.rgb = NAVY
div2_r.font.size = Pt(10)
div2.paragraph_format.space_after = Pt(6)

for line in [
    "報告附件",
    "完整工作流節點文件：嘟嘟房_提案專案 / 99_工作暫存 / DODO_BD_AUTOMATION_PROMPTS.md",
    "工作流品質規則：PROJECT_INDEX.md（Section 15–17）",
    "嘟嘟房完整提案材料：嘟嘟房_提案專案 / 各子資料夾",
]:
    p = doc.add_paragraph()
    r = p.add_run(line)
    r.font.name  = FONT
    r.font.size  = Pt(9.5)
    r.font.color.rgb = DARK_GREY
    r.font.bold  = (line == "報告附件")
    p.paragraph_format.space_after = Pt(2)

# ── 儲存 ─────────────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f"完成：{OUT}")
print(f"大小：{OUT.stat().st_size:,} bytes")
