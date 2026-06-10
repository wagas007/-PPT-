# -*- coding: utf-8 -*-
"""
QA Architecture Reconstruction - Main Executor
"""

import sys
import os
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
sys.path.insert(0, r"C:\Users\C0378\Desktop\旺來PPT製作")

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 載入三份資料 ──────────────────────────────────────────────

# 直接把三份資料 inline（避免 import 問題）
exec(open(r"C:\Users\C0378\Desktop\旺來PPT製作\rebuild_qa_part1.py", encoding="utf-8").read())
QA_ALL = list(QA_DATA)  # from part1

part2_data = {}
exec(open(r"C:\Users\C0378\Desktop\旺來PPT製作\rebuild_qa_part2.py", encoding="utf-8").read(), part2_data)
QA_ALL.extend(part2_data["QA_DATA_2"])

part3_data = {}
exec(open(r"C:\Users\C0378\Desktop\旺來PPT製作\rebuild_qa_part3.py", encoding="utf-8").read(), part3_data)
QA_ALL.extend(part3_data["QA_DATA_3"])

print(f"✓ 共載入 {len(QA_ALL)} 題 Q&A")

# ── 開啟原始文件 ──────────────────────────────────────────────

DOCX_PATH = r"C:\Users\C0378\Desktop\旺來PPT製作\停車場風險報告書.docx"
doc = Document(DOCX_PATH)

# ── 找到 Q&A 章節的起始和結束段落 ────────────────────────────

START_MARKER = "七"       # 第七章開頭
END_MARKER = "八"         # 第八章開頭（Q&A 之後）

start_idx = None
end_idx = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if start_idx is None and text.startswith("七") and "Q&A" in text or \
       start_idx is None and text.startswith("七") and "100" in text or \
       start_idx is None and text.startswith("七") and "決策" in text:
        start_idx = i
        print(f"  找到 Q&A 章節起始：第 {i} 段：「{text[:60]}」")
    elif start_idx is not None and end_idx is None:
        if text.startswith("八") or (text.startswith("八") and "國際" in text) or \
           text.startswith("8.") or text.startswith("八、"):
            end_idx = i
            print(f"  找到 Q&A 章節結束：第 {i} 段：「{text[:60]}」")
            break

# 若未找到精確的標記，用模糊搜索
if start_idx is None:
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "100 題" in text or "100題" in text:
            start_idx = i
            print(f"  模糊匹配 Q&A 起始：第 {i} 段：「{text[:60]}」")
            break

if end_idx is None and start_idx is not None:
    for i, para in enumerate(doc.paragraphs):
        if i <= start_idx:
            continue
        text = para.text.strip()
        if "八、" in text or "八 、" in text or text.startswith("八"):
            end_idx = i
            print(f"  模糊匹配 Q&A 結束：第 {i} 段：「{text[:60]}」")
            break

if start_idx is None:
    print("✗ 找不到 Q&A 章節起始，嘗試在文件末尾插入")
    start_idx = len(doc.paragraphs) - 1
    end_idx = len(doc.paragraphs)

print(f"  Q&A 段落範圍：{start_idx} → {end_idx if end_idx else '文件末尾'}")

# ── 建立新 Q&A 段落的輔助函式 ────────────────────────────────

def add_section_heading(doc, text, level=2):
    """新增章節標題"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # 深藍
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_qa_heading(doc, text):
    """Q 題目標題"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_qa_body(doc, text):
    """Q 答案內文（多行）"""
    for line in text.split("\n"):
        line = line.rstrip()
        if not line:
            # 空行保留段間距
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            continue

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(1)

        # 判斷是否為小標題行（例如「直接回答：」「科學原理：」等）
        is_subhead = False
        subhead_keywords = [
            "直接回答", "原因解釋", "科學原理", "國際案例", "台灣旺來方案",
            "主管機關可能追問", "標準回覆", "風險控制邏輯", "補充說明",
            "物理機制", "具體風險", "管理意涵", "排水溝的雙重風險",
            "不能取代的原因", "三層防護", "四重保護", "三步確認流程",
            "告警分級", "停電的影響", "颱風特有風險", "三道防線",
            "三階段 SOP", "三層防護設計", "六類核保關注資料",
            "退出觸發條件", "財務模型架構", "四步驟", "四個需要",
        ]
        for kw in subhead_keywords:
            if line.startswith(kw) and ("：" in line or "(" in line):
                is_subhead = True
                break

        if is_subhead:
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(10)
        else:
            run = p.add_run(line)
            run.font.size = Pt(10)

    return

# ── 新問題類型分類對照 ────────────────────────────────────────

SECTION_MAP = {
    "A. 物理原理類": ["Q2", "Q3", "Q4", "Q5", "Q7", "Q8", "Q9", "Q40", "Q63"],
    "B. 工程設計類": ["Q14", "Q19", "Q31", "Q32", "Q33", "Q34", "Q35", "Q36",
                     "Q37", "Q38", "Q39", "Q60", "Q61", "Q62"],
    "C. 場址評估類": ["Q11", "Q12", "Q13", "Q15", "Q16", "Q17", "Q18", "Q20",
                     "Q21", "Q22", "Q23", "Q24", "Q25"],
    "D. 容器管理類": ["Q26", "Q27", "Q28", "Q29", "Q30", "Q65"],
    "E. 法規合規類": ["Q56", "Q57", "Q58", "Q59", "Q64", "Q67", "Q68", "Q69", "Q70"],
    "F. 人員操作類": ["Q41", "Q42", "Q43", "Q44", "Q45", "Q46", "Q47", "Q48",
                     "Q49", "Q50", "Q51", "Q52", "Q53", "Q54", "Q55", "Q66", "Q79"],
    "G. 智慧監控類": ["Q6", "Q33", "Q34", "Q35", "Q36", "Q37"],
    "H. 保險責任類": ["Q71", "Q72", "Q73", "Q74", "Q75", "Q77", "Q78", "Q80"],
    "I. 商業試點類": ["Q76", "Q81", "Q82", "Q83", "Q84", "Q85", "Q86",
                     "Q87", "Q88", "Q89", "Q90"],
    "J. 國際比較類": ["Q91", "Q92", "Q93", "Q94", "Q95", "Q96", "Q97",
                     "Q98", "Q99", "Q100"],
    "K. 定義與溝通類": ["Q1", "Q10"],
}

# ── 決定插入位置（在原有 Q&A 章節標題後插入） ──────────────────

# 策略：在 start_idx 段落之後清空到 end_idx，然後插入新內容
# 由於 python-docx 不支援直接刪除段落，我們改用「清空文字 + 替換」的方式
# 實作：找到 Q&A 章節的 XML 父節點，移除舊段落，插入新段落

# 取得 body 的 XML 元素
from lxml import etree

body = doc.element.body
all_paras = body.findall(qn('w:p'))
all_tables = body.findall(qn('w:tbl'))

# 取得所有子元素（段落和表格都是 body 的子元素）
body_children = list(body)

# 找到 start 和 end 對應的 XML 元素
def get_para_xml(doc, idx):
    """取得第 idx 段落的 XML 元素"""
    return doc.paragraphs[idx]._element

# 找到 Q&A 章節標題段落後的所有元素，直到下一章
if start_idx is not None and end_idx is not None:
    start_elem = get_para_xml(doc, start_idx)
    end_elem = get_para_xml(doc, end_idx)

    # 找出 body 中兩個元素之間的所有元素
    found_start = False
    elements_to_remove = []
    for child in list(body):
        if child is start_elem:
            found_start = True
            continue  # 保留章節標題本身
        if found_start:
            if child is end_elem:
                break
            elements_to_remove.append(child)

    print(f"  將移除 {len(elements_to_remove)} 個原有段落/表格")
    for elem in elements_to_remove:
        body.remove(elem)

    # 現在在 end_elem 之前插入新的 Q&A 內容
    # 找到 end_elem 在 body 中的位置
    end_position = list(body).index(end_elem)
    print(f"  新內容將插入在第 {end_position} 個 body 子元素之前")

    # 建立暫時的 document 來生成新段落
    temp_doc = Document()

    # 章節介紹
    p = temp_doc.add_paragraph()
    run = p.add_run("以下 Q&A 每題均區分：直接回答、原因解釋、國際案例、台灣旺來方案、主管機關可能追問、標準回覆。每題由專業角度獨立回答，不使用通用模板。")
    run.font.size = Pt(10)
    run.italic = True
    p.paragraph_format.space_after = Pt(8)

    current_section = None

    for qa in QA_ALL:
        q_label = qa["q"].split(".")[0].strip()  # e.g. "Q1"

        # 確定所屬分類
        qa_section = qa.get("type", "")

        # 如果分類改變，輸出分類標題（依原始 type 欄位）
        if qa_section != current_section:
            current_section = qa_section
            p = temp_doc.add_paragraph()
            run = p.add_run(f"── {qa_section} ──")
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)

        # Q 標題
        p = temp_doc.add_paragraph()
        run = p.add_run(qa["q"])
        run.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)

        # A 答案（多行）
        body_text = qa["a"].strip()
        for line in body_text.split("\n"):
            line_stripped = line.rstrip()
            p = temp_doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(1)

            # 子標題判斷
            is_subhead = False
            for kw in ["直接回答", "原因解釋", "科學原理", "國際案例", "台灣旺來方案",
                       "主管機關可能追問", "標準回覆", "風險形成", "物理機制",
                       "三層防護", "四層防護", "五步驟", "三步驟", "四步驟",
                       "告警分級", "停電的影響", "颱風特有風險", "三道防線",
                       "財務模型架構"]:
                if line_stripped.startswith(kw):
                    is_subhead = True
                    break

            run = p.add_run(line_stripped if line_stripped else " ")
            run.font.size = Pt(10)
            if is_subhead:
                run.bold = True

        # 分隔線
        p = temp_doc.add_paragraph()
        run = p.add_run("─" * 60)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        p.paragraph_format.space_after = Pt(4)

    # 將 temp_doc 的段落元素插入到原始文件的正確位置
    print(f"  正在插入 {len(temp_doc.paragraphs)} 個新段落...")
    insert_before = end_elem

    for para in temp_doc.paragraphs:
        # 複製段落元素到原始文件
        new_para = copy.deepcopy(para._element)
        body.insert(list(body).index(insert_before), new_para)

    print(f"✓ Q&A 章節已更新")

else:
    print("✗ 無法找到 Q&A 章節範圍，跳過注入")

# ── 儲存文件 ──────────────────────────────────────────────────

OUTPUT_PATH = r"C:\Users\C0378\Desktop\旺來PPT製作\停車場風險報告書_QA重建版.docx"
doc.save(OUTPUT_PATH)
print(f"\n✓ 已儲存到：{OUTPUT_PATH}")
print(f"  原始檔案保留於：{DOCX_PATH}")
