# -*- coding: utf-8 -*-
"""
第二輪精修 — 清除殘留問題
"""
import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document

INPUT  = r"C:\Users\C0378\Desktop\旺來PPT製作\停車場風險報告書_QA重建版.docx"
OUTPUT = r"C:\Users\C0378\Desktop\旺來PPT製作\停車場風險報告書_QA重建版.docx"

doc = Document(INPUT)

def replace_para_text(para, new_text):
    from docx.shared import Pt, RGBColor
    if para.runs:
        first_run = para.runs[0]
        bold  = first_run.bold
        size  = first_run.font.size
        try:
            color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.type else None
        except Exception:
            color = None
    else:
        bold, size, color = False, Pt(10), None
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    new_run = para.add_run(new_text)
    new_run.bold = bold
    new_run.font.size = size
    if color:
        try:
            new_run.font.color.rgb = color
        except Exception:
            pass

FIXES = [
    # ── Q48 直接回答：雙括號問題 ──────────────────────────────
    (
        "直接回答：由旺來業務人員（上班時間負責制（10:00–19:00））負責第一時間接收和分類告警；"
        "依告警等級決定是否需要現場人員介入，並在服務時限（SLA）內完成閉環。",

        "直接回答：旺來上班時間（10:00–19:00）由業務人員負責第一時間接收和分類告警；"
        "非上班時間設備依故障安全機制自動保護，場站緊急聯絡人為第一人工窗口，"
        "旺來於次一營業日 10:00 前全面復核所有未閉環告警。"
    ),

    # ── Q48 SLA 說明（Level 1/2/3 的 SLA 是上班時間計算）──────
    (
        "Level 1（一般告警，2 小時 SLA）：",
        "Level 1（一般告警，2 小時 SLA，上班時間計算）："
    ),
    (
        "Level 2（重要告警，30 分鐘 SLA）：",
        "Level 2（重要告警，30 分鐘 SLA，上班時間計算）："
    ),
    (
        "Level 3（緊急告警，即時 SLA）：",
        "Level 3（緊急告警，即時—上班時間立即處理；非上班時間系統自動停用）："
    ),

    # ── Q49 直接回答：殘留「夜間告警與日間同等處理」的舊文本 ──
    (
        "直接回答：旺來採 上班時間監控制度（10:00–19:00）制度，夜間告警與日間同等處理；"
        "遠端停用不受時段限制；Level 3 緊急告警需通知場站緊急聯絡人和消防局，不等到白天。",

        "直接回答：旺來上班時間為 10:00–19:00，非上班時間不設駐守人員。"
        "非上班時間的安全保障依賴三道機制：①系統故障安全自動停用（達告警門檻即鎖定）"
        "②場站緊急聯絡人（設置前預先登記，收到告警簡訊後判斷是否通報 119）"
        "③旺來次一營業日 10:00 前全面復核所有告警。"
        "Level 3 緊急告警（25% LEL、火災）：任何現場人員直接撥打 119，無需等待旺來；"
        "系統已自動停用，消防到場確認安全即可。"
    ),

    # ── Q49 ④ 現場確認時間表：還有舊的「次日清晨 6–8 時」和「1 小時內」 ──
    (
        "·Level 2 夜間告警：次日清晨（6–8 時）安排現場確認\n"
        "·Level 3 夜間告警：視消防介入情況，若消防已到場且確認安全，旺來人員於消防離開後立即前往確認設備狀態；"
        "若無消防介入，旺來人員 1 小時內抵達現場",

        "·Level 2 非上班時間告警：次一營業日 10:00–12:00 完成現場確認\n"
        "·Level 3 非上班時間告警：若消防已介入並確認現場安全，旺來於次一上班日 10:00 前到場確認設備狀態；"
        "若場站緊急聯絡人確認無異味、外觀正常，旺來同樣在次一上班日 10:00–12:00 到場確認"
    ),

    # ── Q49 標準回覆 Q 部分（殘留舊問法）──────────────────────
    (
        "「夜間告警誰負責？有沒有值班制度文件？」",
        "「非上班時間告警誰負責？旺來有沒有 24 小時應變能力？」"
    ),

    # ── 仍有「Level 3 夜間告警」這個用詞的其他段落 ──────────────
    (
        "Level 3 夜間告警",
        "Level 3 非上班時間告警"
    ),
    (
        "Level 2 夜間告警",
        "Level 2 非上班時間告警"
    ),

    # ── Q48 輪班制度段落：若還有「輪班表」 ────────────────────
    (
        "說明業務人員排班與告警處理 SLA；業務人員聯絡方式在與場站合約附件中；可提送稽查。",
        "說明業務人員排班與告警處理 SLA；上班時間聯絡方式和非上班時間場站緊急聯絡人清單在合約附件中；可提送稽查。"
    ),
]

changed = 0
for para in doc.paragraphs:
    original = para.text
    new_text = original
    for old, new in FIXES:
        if old in new_text:
            new_text = new_text.replace(old, new)
    if new_text != original:
        replace_para_text(para, new_text)
        changed += 1

print(f"第二輪精修：{changed} 段落已修改")

# ── 最終全面掃描 ──────────────────────────────────────────────
issues = []
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if "7×24" in t:
        issues.append((i, "7×24", t[:120]))
    if "夜間值班" in t:
        issues.append((i, "夜間值班", t[:120]))
    if "與日間同等處理" in t:
        issues.append((i, "與日間同等處理", t[:120]))
    if "次日清晨（6" in t:
        issues.append((i, "舊時間表", t[:120]))
    if "旺來人員 1 小時內" in t:
        issues.append((i, "舊時間表", t[:120]))

if issues:
    print(f"\n仍有 {len(issues)} 處需確認：")
    for i, tag, txt in issues:
        print(f"  [段落{i}][{tag}] {txt}")
else:
    print("\n全部清除，無殘留問題")

doc.save(OUTPUT)
print(f"\n已儲存：{OUTPUT}")
