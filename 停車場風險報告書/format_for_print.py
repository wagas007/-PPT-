# -*- coding: utf-8 -*-
"""
A4 列印版面調整 — 只改版面，不動內容
"""
import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

INPUT  = r"C:\Users\C0378\Desktop\旺來PPT製作\停車場風險報告書_及100題QA.docx"
OUTPUT = r"C:\Users\C0378\Desktop\旺來PPT製作\停車場風險報告書_A4列印版.docx"

doc = Document(INPUT)

# ── 1. 頁面大小改為 A4，邊距設為專業報告標準 ─────────────────────
A4_W = Cm(21.0)
A4_H = Cm(29.7)

MARGIN_TOP    = Cm(2.5)
MARGIN_BOTTOM = Cm(2.2)
MARGIN_LEFT   = Cm(2.8)   # 左側稍大，預留裝訂空間
MARGIN_RIGHT  = Cm(2.2)

for sec in doc.sections:
    sec.page_width    = A4_W
    sec.page_height   = A4_H
    sec.top_margin    = MARGIN_TOP
    sec.bottom_margin = MARGIN_BOTTOM
    sec.left_margin   = MARGIN_LEFT
    sec.right_margin  = MARGIN_RIGHT
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.4)

print("1. 頁面大小 -> A4 (21.0 x 29.7 cm)")
print("   邊距: 上2.5 下2.2 左2.8 右2.2 cm")

# ── 2. Normal 樣式：行距設 1.15 倍，段後 4pt ──────────────────────
normal_style = doc.styles['Normal']
normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
normal_style.paragraph_format.line_spacing = 1.15
normal_style.paragraph_format.space_after = Pt(4)
normal_style.paragraph_format.space_before = Pt(0)
print("2. Normal 行距 -> 1.15 倍，段後 4pt")

# ── 3. Heading 1：調整間距，讓章節標題更突出 ─────────────────────
h1 = doc.styles['Heading 1']
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after  = Pt(6)
h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
h1.paragraph_format.keep_with_next = True
print("3. Heading 1 -> 段前18pt，段後6pt，不分頁")

# ── 4. Heading 2：調整間距 ────────────────────────────────────────
h2 = doc.styles['Heading 2']
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after  = Pt(4)
h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
h2.paragraph_format.keep_with_next = True
print("4. Heading 2 -> 段前12pt，段後4pt")

# ── 5. List Number 樣式：調整間距 ────────────────────────────────
try:
    ln = doc.styles['List Number']
    ln.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    ln.paragraph_format.line_spacing = 1.15
    ln.paragraph_format.space_after = Pt(3)
    print("5. List Number -> 行距1.15，段後3pt")
except Exception as e:
    print("5. List Number style not found: " + str(e))

# ── 6. 加入頁首（公司名稱 + 文件標題）────────────────────────────
def set_header(section, left_text, right_text):
    header = section.header
    header.is_linked_to_previous = False
    # 清空現有頁首
    for p in header.paragraphs:
        for run in p.runs:
            run.text = ''
    if not header.paragraphs:
        header.add_paragraph()
    hpara = header.paragraphs[0]
    hpara.clear()
    # 加入內容：左側公司名，右側文件名
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    # 左側文字
    run_left = hpara.add_run(left_text)
    run_left.font.size = Pt(8.5)
    run_left.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    # Tab 到右側
    pPr = hpara._p.get_or_add_pPr()
    tabs_elem = OxmlElement('w:tabs')
    tab_elem = OxmlElement('w:tab')
    tab_elem.set(qn('w:val'), 'right')
    # 右側位置：版面寬度 - 左右邊距 = 約 15.0cm = 8505 twips
    tab_elem.set(qn('w:pos'), '8500')
    tabs_elem.append(tab_elem)
    pPr.append(tabs_elem)
    hpara.add_run('\t')
    run_right = hpara.add_run(right_text)
    run_right.font.size = Pt(8.5)
    run_right.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    # 頁首下方細線
    pPr_border = pPr
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr_border.append(pBdr)

# ── 7. 加入頁尾（頁碼）────────────────────────────────────────────
def set_footer_pagenum(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        for run in p.runs:
            run.text = ''
    if not footer.paragraphs:
        footer.add_paragraph()
    fpara = footer.paragraphs[0]
    fpara.clear()
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    fpara.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 頁首上方細線
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = fpara._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top_b = OxmlElement('w:top')
    top_b.set(qn('w:val'), 'single')
    top_b.set(qn('w:sz'), '4')
    top_b.set(qn('w:space'), '1')
    top_b.set(qn('w:color'), 'CCCCCC')
    pBdr.append(top_b)
    pPr.append(pBdr)

    # "- X -" 格式頁碼
    run1 = fpara.add_run('- ')
    run1.font.size = Pt(8.5)
    run1.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # PAGE field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_page = fpara.add_run()
    run_page.font.size = Pt(8.5)
    run_page.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)

    run2 = fpara.add_run(' / ')
    run2.font.size = Pt(8.5)
    run2.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # NUMPAGES field
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instrText2.text = ' NUMPAGES '
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run_total = fpara.add_run()
    run_total.font.size = Pt(8.5)
    run_total.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    run_total._r.append(fldChar3)
    run_total._r.append(instrText2)
    run_total._r.append(fldChar4)

    run3 = fpara.add_run(' -')
    run3.font.size = Pt(8.5)
    run3.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

for i, sec in enumerate(doc.sections):
    set_header(sec,
               '旺來瓦斯股份有限公司',
               '停車場與戶外場域設置風險白皮書')
    set_footer_pagenum(sec)

print("6. 頁首 -> 左:旺來瓦斯，右:文件標題（灰色8.5pt）")
print("7. 頁尾 -> 置中頁碼「- X / N -」（灰色8.5pt）")

# ── 8. Q&A 段落特殊處理：Q 題目標題行（11pt bold）加大段前距 ──────
qa_heading_count = 0
for p in doc.paragraphs:
    if p.runs and p.runs[0].bold and p.runs[0].font.size and p.runs[0].font.size.pt == 11.0:
        txt = p.text.strip()
        if txt.startswith('Q') and ('.' in txt[:5] or len(txt) > 5):
            pf = p.paragraph_format
            if not pf.space_before or pf.space_before.pt < 8:
                pf.space_before = Pt(10)
                pf.space_after = Pt(2)
                pf.keep_with_next = True
                qa_heading_count += 1

print("8. Q&A 題目標題(" + str(qa_heading_count) + "個) -> 段前10pt，keep_with_next")

# ── 9. 分隔線段落（─────）：縮小行距 ────────────────────────────
sep_count = 0
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt.startswith('─') and len(txt) > 10:
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        sep_count += 1

print("9. 分隔線段落(" + str(sep_count) + "個) -> 調整間距")

# ── 10. 儲存 ─────────────────────────────────────────────────────
doc.save(OUTPUT)
print()
print("完成！已儲存：" + OUTPUT)
print()
print("=== 版面調整摘要 ===")
print("  頁面: US Letter -> A4 (21.0 x 29.7 cm)")
print("  上邊距: 1.70 -> 2.50 cm")
print("  下邊距: 1.70 -> 2.20 cm")
print("  左邊距: 1.80 -> 2.80 cm (裝訂側)")
print("  右邊距: 1.80 -> 2.20 cm")
print("  行距: 未設定 -> 1.15 倍")
print("  頁首: 公司名稱 + 文件標題")
print("  頁尾: 置中頁碼 (- X / N -)")
