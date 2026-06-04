"""
產出 Times_FAX_合作提案.docx
傳真封面 + 一頁提案摘要，可直接列印後傳真至 (02) 2507-5667
Run: python generate_fax_docx.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os
from pathlib import Path

OUT = Path(r"c:\Users\C0378\Desktop\旺來PPT製作\PARK24_Times_提案專案\99_工作暫存\Times_FAX_合作提案.docx")

# ── 顏色常數 ──────────────────────────────────────────────────
NAVY    = RGBColor(0, 48, 130)
MID_BLU = RGBColor(0, 70, 160)
DARK_GY = RGBColor(64, 64, 64)
WHITE   = RGBColor(255, 255, 255)
BLACK   = RGBColor(0, 0, 0)
ORANGE  = RGBColor(210, 100, 0)

FILL_NAV  = "003082"
FILL_BLUE = "EBF3FB"
FILL_GREY = "F4F4F4"
FILL_ORNG = "FFF3E0"


# ── 輔助函式 ──────────────────────────────────────────────────
def set_cell_bg(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def set_cell_padding(cell, top=60, bottom=60, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def no_border_table(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)

def thin_border_table(table, color="999999"):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)

def set_para_border_bottom(para, color_hex="003082", sz=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def p_run(doc, text, size=11, bold=False, italic=False, color=None, font="Calibri", align=None):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    return p

def p_section(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    set_para_border_bottom(p)
    return p

def p_spacer(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(" ")
    run.font.size = Pt(pt)
    return p

def p_callout(doc, text, fill=FILL_BLUE, color=NAVY, bold=True):
    table = doc.add_table(rows=1, cols=1)
    no_border_table(table)
    cell = table.cell(0, 0)
    set_cell_bg(cell, fill)
    set_cell_padding(cell, top=100, bottom=100, left=160, right=160)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cp.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = bold
    run.font.color.rgb = color
    return table

def p_info_table(doc, rows_data):
    """兩欄資訊表：標籤 | 值，無邊框"""
    table = doc.add_table(rows=len(rows_data), cols=2)
    no_border_table(table)
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(13)
    for i, (label, value) in enumerate(rows_data):
        c0, c1 = table.cell(i, 0), table.cell(i, 1)
        set_cell_padding(c0, top=50, bottom=50, left=80, right=80)
        set_cell_padding(c1, top=50, bottom=50, left=80, right=80)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.name = "Calibri"; r0.font.size = Pt(10.5); r0.bold = True
        r0.font.color.rgb = DARK_GY
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(value)
        r1.font.name = "Calibri"; r1.font.size = Pt(10.5)
        r1.font.color.rgb = BLACK
    return table

def add_page_break(doc):
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)

def p_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run_bullet = p.add_run("• ")
    run_bullet.font.name = "Calibri"
    run_bullet.font.size = Pt(10.5)
    run_bullet.font.color.rgb = NAVY
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.font.name = "Calibri"; rb.font.size = Pt(10.5); rb.bold = True
        rb.font.color.rgb = BLACK
        run = p.add_run(text)
    else:
        run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    return p


# ── 主程式 ──────────────────────────────────────────────────
doc = Document()

# 頁面設定：A4，左右 2.8cm，上 2.5cm，下 2.0cm
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.8)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ═══════════════════════════════════════════════════════════
# PAGE 1：傳真封面頁
# ═══════════════════════════════════════════════════════════

# 大標題：傳  真
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("傳　　真")
run.font.name = "Calibri"
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = NAVY
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(4)
set_para_border_bottom(p, color_hex="003082", sz=12)

p_spacer(doc, 8)

# 收發件資訊表
p_info_table(doc, [
    ("收件人：", "方先生"),
    ("收件公司：", "台灣普客二四股份有限公司　Times 停車場"),
    ("傳真號碼：", "(02) 2507-5667"),
    ("", ""),
    ("發件人：", "_______________（請填入您的姓名）"),
    ("發件部門：", "旺來瓦斯股份有限公司　開發部"),
    ("聯絡電話：", "_______________（請填入您的電話）"),
    ("傳真日期：", "2026 年 06 月 03 日"),
    ("總頁數：", "共 2 頁（含封面）"),
])

p_spacer(doc, 10)

# 主旨欄
p_section(doc, "主旨")
p_run(doc, "停車場場站服務節點合作提案洽詢", size=11.5, bold=True, color=BLACK)

p_spacer(doc, 8)

# 封面內文
p_section(doc, "致　方先生")
p_spacer(doc, 4)

body_text = [
    "方先生您好，",
    "",
    "感謝您今日撥冗接聽電話，並惠予指引傳真聯繫方式。",
    "",
    "本公司旺來瓦斯股份有限公司開發部，謹以此傳真呈送「停車場場站服務節點合作提案摘要」（第 2 頁），",
    "敬請惠予轉交　貴公司業務開發或場站合作相關負責人。",
    "",
    "本提案核心主張如下：",
]

for line in body_text:
    if line == "":
        p_spacer(doc, 4)
    else:
        p_run(doc, line, size=10.5, color=BLACK)

p_spacer(doc, 4)
p_callout(
    doc,
    "旺來瓦斯希望與 PARK24 / Times 先以 1 至 3 個示範場站驗證「社區型智慧服務節點」，\n"
    "旺來負責全部設備、補貨、維運與客服，Times 只需協助現勘窗口安排。\n"
    "試點期間 90 天，以數據決定是否繼續擴點。",
    fill=FILL_BLUE, color=NAVY
)
p_spacer(doc, 8)

closing_lines = [
    "若　貴公司有任何問題，歡迎來電洽詢。",
    "如需安排 30 分鐘初步說明會議，旺來可配合　貴公司時間。",
    "",
    "敬請 惠覽，謝謝。",
    "",
    "旺來瓦斯股份有限公司　開發部　敬上",
    "_______________（請填入您的姓名）",
    "電話：_______________",
]
for line in closing_lines:
    if line == "":
        p_spacer(doc, 4)
    else:
        p_run(doc, line, size=10.5, color=BLACK)

# ═══════════════════════════════════════════════════════════
# PAGE 2：一頁提案摘要
# ═══════════════════════════════════════════════════════════
add_page_break(doc)

# 頁眉標題
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("旺來瓦斯 × PARK24 / Times　場站合作提案摘要")
run.font.name = "Calibri"
run.font.size = Pt(16)
run.bold = True
run.font.color.rgb = NAVY
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(2)
set_para_border_bottom(p, color_hex="003082", sz=10)

p_run(doc, "本文件為初步提案摘要，詳細方案待雙方安排 30 分鐘說明會議後再行呈報。",
      size=9.5, italic=True, color=DARK_GY, align="center")
p_spacer(doc, 8)

# Section 1: 合作主張
p_section(doc, "一、合作主張")
p_callout(
    doc,
    "旺來瓦斯不是來承租停車格，而是協助 Times 在場站邊角空間建立「社區型智慧服務據點」，\n"
    "活化閒置空間，不影響停車本業，由旺來負責全部設備與維運。",
    fill=FILL_BLUE, color=NAVY
)
p_spacer(doc, 6)

# Section 2: 旺來提供什麼
p_section(doc, "二、旺來負責的全部事項（Times 無需額外投入）")
p_bullet(doc, "設備購置、安裝施工與法規申請",           bold_prefix="設備：")
p_bullet(doc, "商品補貨、日常作業與庫存管理",           bold_prefix="補貨：")
p_bullet(doc, "消費者客服與異常排除（24 小時）",         bold_prefix="客服：")
p_bullet(doc, "設備維護、責任保險與安全文件",           bold_prefix="維運：")
p_bullet(doc, "試點 90 天後提供完整消費者行為數據報告", bold_prefix="數據：")
p_spacer(doc, 6)

# Section 3: Times 只需要做什麼
p_section(doc, "三、Times 只需要做三件事")
p_bullet(doc, "提供桃園或台中 1 至 3 個候選場站基本資料",   bold_prefix="Step 1：")
p_bullet(doc, "安排一次現勘（每場站約 30 分鐘，旺來自行前往）", bold_prefix="Step 2：")
p_bullet(doc, "確認場站管理人員基本配合意願",                  bold_prefix="Step 3：")
p_spacer(doc, 6)

# Section 4: 試點框架
p_section(doc, "四、試點框架")
# 四格表
table = doc.add_table(rows=2, cols=2)
thin_border_table(table, color="CCCCCC")
labels = [
    ("試點規模", "1 至 3 個示範場站（桃園或台中優先）"),
    ("試點期間", "90 天，以 KPI 數據決定是否擴點"),
    ("Time to Start", "雙方確認試點場站後，設備安裝 30 天內完成"),
    ("Times 風險", "零資本投入、無現場管理負擔、可隨時回顧數據"),
]
cells = [table.cell(r, c) for r in range(2) for c in range(2)]
for cell, (label, value) in zip(cells, labels):
    set_cell_bg(cell, FILL_GREY)
    set_cell_padding(cell, top=80, bottom=80, left=120, right=120)
    cp = cell.paragraphs[0]
    rb = cp.add_run(label + "　")
    rb.bold = True; rb.font.name = "Calibri"; rb.font.size = Pt(10)
    rb.font.color.rgb = NAVY
    rv = cp.add_run(value)
    rv.font.name = "Calibri"; rv.font.size = Pt(10)
    rv.font.color.rgb = BLACK

p_spacer(doc, 8)

# Section 5: 下一步
p_section(doc, "五、唯一的請求")
p_callout(
    doc,
    "請協助安排 30 分鐘總部說明會議，或提供業務開發窗口的聯絡方式。\n"
    "會議目標：確認試點意願、候選場站方向、現勘流程安排。\n"
    "旺來可配合貴公司任何時間，並可親至貴公司辦公室簡報說明。",
    fill=FILL_ORNG, color=ORANGE, bold=True
)

p_spacer(doc, 10)
p_run(doc, "旺來瓦斯股份有限公司　開發部　_______________", size=9.5, color=DARK_GY)
p_run(doc, "電話：_______________　　Email：_______________", size=9.5, color=DARK_GY)


doc.save(str(OUT))
print(f"Done: {OUT.name}  {os.path.getsize(OUT):,} bytes")
