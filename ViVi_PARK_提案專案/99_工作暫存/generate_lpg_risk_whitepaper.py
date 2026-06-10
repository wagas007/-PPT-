from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TARGET = next(p for p in ROOT.glob("*.docx") if p.name.startswith("停車場風險報告書"))
BACKUP = ROOT / "停車場風險報告書_原始備份.docx"
IMAGE = ROOT / "澎湖金龍發智取櫃_北屯智取櫃.png"

ORANGE = "F05A24"
DARK = "202020"
LIGHT = "FFF4EE"
PALE = "FFF9F6"
GRAY = "666666"
RED = "C00000"
GREEN = "168A4A"


SOURCES = [
    ("TW-01", "台灣消防署", "公共危險物品及可燃性高壓氣體製造儲存處理場所設置標準暨安全管理辦法", "https://www.nfa.gov.tw/upload/pro/attachment/d3f1e8d2b0e71fef0fbcdab079391917.pdf"),
    ("TW-02", "台灣勞動部職安署", "液化石油氣容器串接氣體供應裝置使用作業指引", "https://www.osha.gov.tw/48110/48713/48735/137104/"),
    ("TW-03", "台灣消防署", "容器保管室制度及液化石油氣安全管理對策", "https://www.nfa.gov.tw/cht/index.php?code=list&ids=1276"),
    ("US-01", "美國 OSHA", "29 CFR 1910.110 Storage and handling of liquefied petroleum gases", "https://www.osha.gov/laws-regs/regulations/standardnumber/1910/1910.110"),
    ("UK-01", "英國 HSE", "LPG issues: outdoor cylinder storage and DSEAR", "https://www.hse.gov.uk/electricity/atex/lpg.htm"),
    ("UK-02", "英國 HSE", "Drum / cylinder handling technical measures", "https://www.hse.gov.uk/comah/sragtech/techmeascylinder.htm"),
    ("UK-03", "英國 HSE", "Bulk LPG storage tank: siting, ventilation, security and impact protection", "https://www.hse.gov.uk/gas/lpg/storagetank.htm"),
    ("DE-01", "德國 BAuA", "TRGS 510 Storage of hazardous substances in non-stationary containers", "https://www.baua.de/DE/Angebote/Regelwerk/TRGS/pdf/TRGS-510.html?blob=publicationFile&v=6"),
    ("FR-01", "法國 Légifrance", "Article GZ 7: commercial propane cylinder storage", "https://www.legifrance.gouv.fr/codes/section_lc/JORFTEXT000000290033/LEGISCTA000020304213/2019-07-01"),
    ("FR-02", "法國 INRS", "Les bouteilles de gaz, ED 6369", "https://www.inrs.fr/dam/inrs/CataloguePapier/ED/TI-ED-6369.pdf"),
    ("NZ-01", "紐西蘭 EPA", "LPG at home: safe outdoor storage and quantity controls", "https://www.epa.govt.nz/everyday-environment/using-lpg-gas-at-home/"),
    ("NZ-02", "紐西蘭 EPA", "LPG Compliance 100 kg to 300 kg", "https://www.epa.govt.nz/assets/Uploads/Documents/Hazardous-Substances/Guidance/HSNOCOP-50-LPG-Compliance-100kg-to-300kg.pdf"),
    ("AU-01", "澳洲 WorkSafe Queensland", "Gases in cylinders", "https://www.worksafe.qld.gov.au/safety-and-prevention/hazards/hazardous-chemicals/specific-hazardous-chemicals/gases-in-cylinders"),
    ("AU-02", "澳洲 WorkSafe Queensland", "Handling and storing LPG cylinders / AS/NZS 1596 reference", "https://www.worksafe.qld.gov.au/__data/assets/pdf_file/0019/16732/attach-lpg-cylinder.pdf"),
    ("JP-01", "日本經濟產業省 METI", "LP ガスの安全・規制", "https://www.meti.go.jp/policy/safety_security/industrial_safety/sangyo/lpgas/anzen_torikumi/"),
    ("JP-02", "日本經濟產業省 METI", "LP ガス事故情報", "https://www.meti.go.jp/policy/safety_security/industrial_safety/sangyo/lpgas/lpjiko/index.html"),
    ("JP-03", "日本經濟產業省 METI", "液化石油ガス充てん容器の流出防止対策", "https://www.safety-chubu.meti.go.jp/sangyou/lpgas/information_2030_01.html"),
    ("KR-01", "韓國國家法令資訊中心", "액화석유가스의 안전관리 및 사업법 시행규칙", "https://www.law.go.kr/LSW/lsPdfPrint.do?ancYnChk=0&bylChaChk=N&efGubun=Y&efYd=20231010&joAllCheck=Y&joEfOutPutYn=on&lsiSeq=255297"),
]


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell(cell, text, bold=False, color=None, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Microsoft JhengHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows, header_color=ORANGE, size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        shade(t.rows[0].cells[i], header_color)
        set_cell(t.rows[0].cells[i], h, True, "FFFFFF", size)
    for idx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value, False, None, size)
            if idx % 2:
                shade(cells[i], PALE)
    doc.add_paragraph()
    return t


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)


def manual_numbered(doc, items):
    for index, text in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.65)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(f"{index}. ")
        p.add_run(text)


def labeled_paragraph(doc, label, text, bold=False):
    p = doc.add_paragraph()
    label_run = p.add_run(label)
    label_run.bold = bold
    p.add_run(text)
    return p


def callout(doc, title, text, color=LIGHT):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.cell(0, 0)
    shade(c, color)
    p = c.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(ORANGE)
    p2 = c.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


def add_header_footer(section):
    header = section.header.paragraphs[0]
    header.text = "旺來瓦斯｜LPG 智取櫃停車場與戶外場域設置風險白皮書"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("誠懇・專業・當責・共榮　｜　本文件供決策與初步審查使用，不取代主管機關、消防、工程或保險正式審查")
    footer.runs[0].font.size = Pt(7.5)
    footer.runs[0].font.color.rgb = RGBColor.from_string(GRAY)


if not BACKUP.exists():
    copy2(TARGET, BACKUP)

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.7)
sec.bottom_margin = Cm(1.7)
sec.left_margin = Cm(1.8)
sec.right_margin = Cm(1.8)
add_header_footer(sec)

styles = doc.styles
styles["Normal"].font.name = "Microsoft JhengHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
styles["Normal"].font.size = Pt(10)
styles["Normal"].paragraph_format.space_after = Pt(5)
for name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
    styles[name].font.name = "Microsoft JhengHei"
    styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
styles["Heading 1"].font.color.rgb = RGBColor.from_string(ORANGE)
styles["Heading 2"].font.color.rgb = RGBColor.from_string(DARK)
styles["Heading 3"].font.color.rgb = RGBColor.from_string(ORANGE)

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("台灣桶裝瓦斯智取櫃\n停車場與戶外場域設置風險白皮書")
r.bold = True
r.font.size = Pt(25)
r.font.color.rgb = RGBColor.from_string(ORANGE)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Decision-Oriented Risk Assessment White Paper")
r2.italic = True
r2.font.size = Pt(14)
if IMAGE.exists():
    pic = doc.add_picture(str(IMAGE), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
callout(
    doc,
    "白皮書核心結論",
    "LPG 智取櫃並非普通販賣機，也不應被視為零風險設備；它是可透過場址篩選、戶外通風、合格容器、限制接觸、防撞、監控、巡檢、責任分工與保險，進行「一場一審」的受控能源服務節點。台灣實際設置前，仍須由轄區消防機關確認法律定性、允許儲放量、場所要求與申請程序。",
)
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.add_run("旺來瓦斯股份有限公司｜2026 年 6 月｜外部合作夥伴決策版").bold = True

page_break(doc)
doc.add_heading("文件定位與使用邊界", level=1)
doc.add_paragraph("本文件供總經理、停車場業者、消防與經發主管機關、加盟商、保險公司及場域管理人進行初步決策與風險溝通。")
table(doc, ["本文件可以做什麼", "本文件不能取代什麼"], [
    ["建立共同風險語言、初步場址篩選、風險矩陣、控制措施、Q&A 與試點 Gate。", "消防審查、建築／電氣設計簽證、容器檢驗、保險核保、主管機關核准與個案現勘。"],
    ["比較國際最佳實務並轉譯成台灣可執行方案。", "將國外法規直接視為台灣法定標準。"],
    ["協助合作夥伴判斷是否值得進入下一階段審查。", "承諾任何場站必然可設置，或宣稱零風險。"],
])
doc.add_heading("法規定性先行原則", level=2)
doc.add_paragraph("智慧櫃同時涉及 LPG 容器、販賣／交付、暫存、場址與無人化取用。台灣現行法規可能依實際儲放量、容器狀態、營運模式及場址，適用不同要求。因此，任何投資與施工前，應以書面向轄區消防機關確認下列事項：")
for x in ["智慧櫃在個案中屬販賣場所、容器保管室、儲存場所、處理場所，或其他管理類型。", "灌氣容器、殘氣容器與交換後空瓶的計量及分區方式。", "允許儲放量、必要安全距離、消防設備、漏氣警報、通風與申請文件。", "無人化取用、補貨、遠端監控與緊急應變的管理責任。"]:
    numbered(doc, x)

doc.add_heading("主管決策摘要：五個問題", level=1)
table(doc, ["問題", "決策結論"], [
    ["1. 能不能做？", "有條件可評估。優先戶外或高度通風、可防撞、可監控、可控量、權責清楚的場址；設置前須完成主管機關確認。"],
    ["2. 最大風險是什麼？", "不是單一菸蒂，而是「容器或閥件失效／操作異常 + 氣體累積 + 有效點火源 + 未及時發現處置」的組合。"],
    ["3. 最大誤解是什麼？", "誤以為只要有 LPG 就必然爆炸，或反過來誤以為櫃體與複合材料容器可消除所有風險。"],
    ["4. 最大成功因素是什麼？", "法律定性先行、一場一審、低量試點、戶外通風、防撞、限制接觸、監控、巡檢、保險與退出機制。"],
    ["5. 旺來如何落地？", "先完成文件審查，再選 A 級場址進行限量試點；以 90 天數據決定是否擴大。"],
])

page_break(doc)
doc.add_heading("一、國際最佳實務：國外為何能運作", level=1)
doc.add_paragraph("本研究未發現「所有國家普遍允許任意設置全自動 LPG 販賣機」的證據。較成熟、可驗證的國際模式，是零售通路、加油站、五金店或戶外場域設置受控的 cylinder exchange cage／outdoor cylinder store。自動化可改善權限與紀錄，但不會取消 LPG 儲存與場址安全義務。")
table(doc, ["國家／地區", "官方規範或實務重點", "對旺來的啟示", "來源"], [
    ["美國", "OSHA 1910.110 要求容器避免高溫、物理損壞與未授權干擾；多數容器原則上位於建築外。", "櫃體位置、防撞、限制接觸與容器保護是核心。", "US-01"],
    ["加拿大", "依省級採用之 CSA B149 系列管理；臨時設置亦要求防高溫、換瓶後檢漏。", "每次補貨／換桶後檢核與設備核准不可省略。", "加拿大省級技術安全規範"],
    ["英國", "HSE 偏好戶外儲存以利氣體散逸；要求場址風險評估、安全距離、保全、防撞與緊急安排。", "戶外優先；無法符合戶外儲存條件時，需更嚴格分區與設備要求。", "UK-01/02/03"],
    ["德國", "TRGS 510 以危害評估、存取限制、儲存組織、警報、檢查、消防與加壓氣體專章管理。", "智慧化應用在存取權限、庫存紀錄、警報與檢查，而非弱化標準。", "DE-01"],
    ["法國", "公眾建築的商用 propane 容器外部儲存應避免公眾可自由接觸，不阻礙人車通行；需通風、防火與整潔。", "停車場必須避開主要人車路徑，並限制未授權接觸。", "FR-01/02"],
    ["荷蘭／挪威", "多依歐洲危險物質、爆炸性環境與國家消防規範進行風險分區及場址管理。", "台灣不應直接照搬距離；可採相同的風險分區與權責精神。", "歐盟 ATEX／各國規範"],
    ["澳洲", "以 AS/NZS 1596 與職安機關指引管理 LPG 儲存與搬運，重視容器安全、搬運與危害化學品管理。", "補貨與搬運事故往往比重大火災更常見，SOP 與人員訓練很重要。", "AU-01/02"],
    ["紐西蘭", "戶外、直立、乾燥通風、堅固耐火基座、防傾倒；依儲量提高證明、通報及認證要求。", "採低量試點並設定儲量升級 Gate；排水溝與低窪處需避開。", "NZ-01/02"],
    ["日本", "供應商負有設備檢查與保安責任；事故分類包含洩漏、洩漏火災／爆炸、容器遺失與竊盜；水災區需防容器流失。", "納入防災、竊盜、庫存與容器固定；不是只看火災。", "JP-01/02/03"],
    ["韓國", "以液化石油氣安全管理及事業法規範儲存設備、容器保管與供應商安全檢查。", "營運者與供應者持續檢查責任應寫入合約。", "KR-01"],
], size=7.7)

doc.add_heading("國際共通的五層風險管理", level=2)
table(doc, ["層次", "國際共通做法", "旺來台灣版"], [
    ["技術措施", "戶外通風、容器直立固定、閥件保護、防撞、警報、適當電氣設備。", "櫃體通風、容器獨立定位、CCTV、門狀態、溫度／漏氣監測、防撞設施。"],
    ["管理措施", "限制未授權接觸、庫存控制、定期巡檢、異常隔離、標示。", "APP 權限、一次一格、補貨人員認證、每日遠端檢查、定期現場巡檢。"],
    ["法規措施", "依儲量、場址與用途分級；高風險場所需許可、證明或檢查。", "消防機關書面確認、場址文件、容器檢驗與必要申請。"],
    ["保險措施", "依場址、庫存、第三人暴露與營運責任核保。", "商品責任、公共意外、設備財損、第三人與營業中斷責任界定。"],
    ["營運措施", "供應商承擔補貨、檢查、緊急應變與紀錄。", "旺來負責補貨、客服、異常桶隔離、通報、回收與資料留存。"],
])

page_break(doc)
doc.add_heading("二、MECE 風險樹與事故頻率排序", level=1)
table(doc, ["風險類別", "主要風險", "關鍵控制"], [
    ["A 場址", "封閉／低窪、靠近火源、車道撞擊、排水溝、淹水、逃生阻礙。", "戶外優先、場址 Gate、排水與淹水評估、防撞、禁設清單。"],
    ["B 設備", "櫃門、鎖控、電氣、感測器、通風、固定基礎失效。", "工程驗收、故障安全設計、告警、定期保養、斷電應變。"],
    ["C 容器", "過期、損傷、閥件洩漏、翻倒、殘氣桶混放。", "合格容器、掃碼追溯、直立固定、滿空分區、異常桶隔離。"],
    ["D 人員", "誤操作、未授權取用、補貨搬運傷害、吸菸。", "APP 身分、操作引導、補貨訓練、禁菸、客服與遠端介入。"],
    ["E 營運", "補貨錯誤、巡檢漏失、告警未處理、庫存不符。", "SOP、服務時限（SLA）、稽核、事件分級、庫存與告警紀錄。"],
    ["F 法規", "定性錯誤、超量、未申請、文件不完整。", "消防預審、法規清單、變更管理、定期合規稽核。"],
    ["G 品牌", "民眾疑慮、社群事件、合作夥伴聲譽受損。", "透明文件、客服、危機溝通、試點與退出機制。"],
    ["H 財務", "低使用率、保費／施工過高、事故損失、停業。", "單場 Business Case、限量試點、保險、明確停損。"],
])

doc.add_heading("事故頻率排序：先管常見事件，再防重大事故", level=2)
doc.add_paragraph("公開資料不足以支持台灣智慧櫃的精確事故率。以下排序是依國際事故分類、一般戶外設備暴露與營運經驗建立的「初步頻率假設」，須以試點事件資料校正。")
table(doc, ["頻率假設", "事件", "為何重要", "管理優先度"], [
    ["高頻／低至中影響", "誤操作、門鎖／掃碼失敗、補貨延誤、客訴、容器放置錯誤、監控斷線。", "最常影響服務與品牌，容易累積成管理漏洞。", "最高"],
    ["中頻／中至高影響", "車輛擦撞、惡意破壞、偷竊、容器翻倒、積水／淹水、高溫、短暫漏氣告警。", "停車場暴露的主要差異風險。", "高"],
    ["低頻／高影響", "持續洩漏、火災、爆炸、重大人身傷害、颱風造成櫃體或容器位移。", "頻率低但不可接受，需多層防護與緊急應變。", "最高／零容忍"],
])

doc.add_heading("初步風險矩陣", level=2)
risk_rows = [
    ["車輛撞擊", "中", "重大", "高", "位置退縮、防撞柱、視線與倒車熱區審查", "中低"],
    ["未授權接觸／破壞", "中", "中高", "中高", "鎖控、APP 權限、CCTV、告警", "低中"],
    ["容器翻倒／閥件損傷", "中", "重大", "高", "直立定位、閥件保護、一次一格、補貨檢查", "低中"],
    ["漏氣未及時處置", "低中", "重大", "高", "合格容器、漏氣偵測、通風、告警與隔離 SOP", "低中"],
    ["菸蒂／點火源", "中", "重大", "高", "禁菸、通風、避免氣體累積、監控與巡檢", "中低"],
    ["高溫曝曬", "中", "中高", "中高", "遮陽、溫度監測、40°C 管理門檻、停用機制", "低中"],
    ["淹水／颱風", "低中", "重大", "高", "淹水圖資、固定、防流失、颱風前停用／撤桶", "低中"],
    ["法規定性錯誤", "中", "重大", "高", "消防書面預審、設置前 Gate", "低"],
    ["品牌客訴", "中高", "中高", "高", "透明說明、合作夥伴審稿、客服與事件通報", "中"],
]
table(doc, ["風險", "可能性", "影響", "初始風險", "控制措施", "殘餘風險"], risk_rows, size=8)

page_break(doc)
doc.add_heading("三、停車場適設性與禁止條件", level=1)
callout(doc, "決策原則", "停車場不是一律適合，也不是一律禁止。適合與否取決於是否能將 LPG 節點與車流、低窪積氣、火源、公眾自由接觸及逃生動線有效隔離。初期優先戶外平面場、半戶外且高度通風場；地下或封閉停車場原則上不列入首波試點。")
table(doc, ["優先考慮", "條件式考慮", "禁止／退件條件"], [
    ["戶外平面、自然通風、邊角清楚、可設防撞、CCTV 覆蓋、補貨可安全短停、權責單純。", "半戶外、鄰近建物／排水設施、第三方場域、靠近 EV 設備；須完成工程與主管機關審查。", "地下／明顯封閉低窪、逃生通道、車道轉角或倒車熱區、無法防撞、常態火源、淹水高風險、無合法權責或消防不接受。"],
])
doc.add_heading("停車場必須回答的十個 Gate", level=2)
manual_numbered(doc, [
    "法律定性與允許儲放量是否已由轄區消防機關確認？",
    "是否為戶外或足以避免 LPG 累積的通風環境？",
    "是否避開地下室、低窪處、排水溝、集水井與氣體可能積聚區？",
    "是否避開車道轉角、倒車熱區、出入口、逃生路徑與主要行人動線？",
    "是否能設置可靠防撞、固定基礎與容器防傾倒措施？",
    "是否能維持台灣要求之禁菸火、溫度與消防條件？",
    "CCTV、門狀態、溫度、漏氣與通訊告警是否可用？",
    "補貨、異常桶隔離、停用、撤桶及緊急通報是否可執行？",
    "地主、停車場、旺來、保險與主管機關權責是否書面清楚？",
    "若發生告警、客訴、颱風或法規疑義，是否能立即停用與退出？",
])

doc.add_heading("四、100 分場址評分表", level=1)
doc.add_paragraph("本評分表是旺來內部／合作夥伴的初步篩選工具，不是法規核准。任何法定禁止或消防不接受條件，均可直接否決，不因總分高而通過。")
score_rows = [
    ["法規與權責", 20, "消防初步確認、土地／場站同意、用途與責任清楚"],
    ["通風與積氣風險", 15, "戶外自然通風；無地下、低窪、集水或氣體滯留疑慮"],
    ["車道與防撞", 15, "避開倒車熱區／轉角；可設防撞與退縮"],
    ["火源與電氣", 10, "符合台灣禁菸火要求；電氣與充電設備經專業審查"],
    ["建物、人員與逃生", 10, "不阻礙門窗、逃生、人車通道；公眾暴露可控"],
    ["排水、淹水與天候", 10, "遠離排水積氣點；低淹水／颱風風險，可固定與撤桶"],
    ["監控與通訊", 8, "CCTV、門、溫度、漏氣、斷線與告警可監控"],
    ["補貨與緊急應變", 7, "安全補貨、異常桶隔離、消防到達與撤離條件"],
    ["品牌與社區接受度", 5, "外觀、告示、客服與鄰里溝通可控"],
]
table(doc, ["評分構面", "配分", "滿分條件"], score_rows)
table(doc, ["等級", "分數", "決策"], [
    ["A 級", "85–100", "可進入消防／工程／保險正式初審與低量試點設計。"],
    ["B 級", "70–84", "需完成指定改善後再審；不得直接施工。"],
    ["C 級", "未滿 70", "不建議設置。"],
    ["否決項", "不論分數", "地下／積氣風險、無法防撞、阻礙逃生、權責不明、消防不接受、無緊急退出能力。"],
])

page_break(doc)
doc.add_heading("五、旺來台灣版落地方案", level=1)
table(doc, ["控制層", "旺來建議方案", "驗證／紀錄"], [
    ["容器", "僅使用檢驗有效、外觀與閥件合格容器；掃碼追溯；滿桶、殘氣桶與異常桶分流。", "容器 ID、檢驗期限、補貨前後照片、異常處理紀錄"],
    ["櫃體", "不燃／難燃思維、自然通風、一次一格、直立固定、閥件保護、防撬與固定基礎。", "設計圖、材料規格、工程驗收"],
    ["場址", "戶外優先、避開低窪排水與火源、人車分流、防撞、遮陽、防淹水。", "現勘表、配置圖、照片、消防意見"],
    ["智慧監控", "APP 權限、一次性取用碼、櫃門狀態、CCTV、溫度、漏氣、斷線與電力告警。", "告警紀錄、存取紀錄、權限稽核"],
    ["營運", "補貨雙重檢查、日常遠端監看、定期巡檢、異常桶隔離、颱風與高溫停用 SOP。", "SOP、巡檢表、教育訓練、事件報告"],
    ["應變", "告警分級、遠端停用、現場隔離、通知消防／場站、撤桶與事故調查。", "通報樹、演練、服務時限（SLA）、改善追蹤"],
    ["保險與合約", "商品責任、公共意外、設備財損、第三人責任；明確界定場站與旺來責任。", "保單、合約、責任矩陣"],
])
doc.add_heading("90 天低量試點 Gate", level=2)
table(doc, ["階段", "必須完成", "通過條件"], [
    ["Gate 0：法律確認", "消防書面預審、儲放量與場所定性、場站權責。", "無重大法規障礙。"],
    ["Gate 1：設計審查", "場址評分 A 級、配置、防撞、通風、監控、保險。", "跨部門簽核。"],
    ["Gate 2：安裝驗收", "工程、告警、停用、監控、補貨及應變測試。", "缺失關閉後上線。"],
    ["Gate 3：低量營運", "限量、限時、遠端監控、加密巡檢。", "無重大告警；異常均在約定服務時限（SLA）內處理。"],
    ["Gate 4：90 天回顧", "使用量、告警、客訴、撞擊、異常桶、維運、成本。", "殘餘風險可接受，才考慮下一場。"],
])

doc.add_heading("六、主管簡報版：十頁結論", level=1)
for x in [
    "P1｜LPG 智取櫃可被審查，但不能被當作普通販賣機。",
    "P2｜國際成熟模式靠戶外通風、限制接觸、防撞、巡檢與供應商責任運作。",
    "P3｜最大風險是多條件同時成立，不是單一菸蒂。",
    "P4｜高頻問題多為操作、補貨、撞擊、破壞與告警；重大火災低頻但零容忍。",
    "P5｜停車場是否適合，取決於通風、車流、低窪積氣、防撞、權責與應變。",
    "P6｜地下、封閉、低窪、逃生路徑與無法防撞場址不列首波。",
    "P7｜100 分評分表只能篩選，法定否決項仍優先。",
    "P8｜旺來以容器、櫃體、監控、營運、應變與保險形成多層防護。",
    "P9｜先做一個 A 級場址、低量、可退出的 90 天試點。",
    "P10｜每新增一場，均重新完成消防、工程、保險與場站 Business Case。",
]:
    bullet(doc, x)

# Stakeholder briefs
page_break(doc)
doc.add_heading("七、利害關係人說明版", level=1)
doc.add_heading("消防局說明版", level=2)
doc.add_paragraph("旺來希望在施工前確認智慧櫃之法律定性、允許儲放量與必要安全要求。本案採戶外優先、一場一審、限量試點，並提供容器、櫃體、通風、防撞、監控、消防、補貨、異常與保險資料；任何主管機關不接受或條件無法滿足之場址均不設置。")
doc.add_heading("停車場業者說明版", level=2)
doc.add_paragraph("停車場不需承擔 LPG 商品、補貨、客服與設備營運責任；旺來承接日常營運及異常處理。場站保留位置、品牌、動線與停用決定權。設置前逐場確認通風、車流、防撞、監控、保險與退出機制。")
doc.add_heading("加盟商說明版", level=2)
doc.add_paragraph("加盟商不得自行選點、增量或改裝。補貨、巡檢、異常桶隔離、通報與紀錄均依旺來 SOP 執行；任何告警、容器損傷、淹水、高溫或設備異常，均先停用、隔離並通報。")

# 100 QA
page_break(doc)
doc.add_heading("八、100 題決策 Q&A", level=1)
doc.add_paragraph("以下 Q&A 採外部合作夥伴可直接閱讀的保守表述。每題均區分專業回答、國際最佳實務與台灣落地方式。")

qa_topics = {
    "基本原理與風險認知": [
        "LPG 智取櫃是不是普通販賣機？", "LPG 有沒有可能爆炸？", "菸蒂是否會直接引爆？", "空桶是不是就沒有風險？",
        "複合材料容器是否代表零風險？", "智慧監控是否能取代現場巡檢？", "為什麼優先設在戶外？", "為什麼 LPG 洩漏特別注意低窪處？",
        "火災和爆炸哪一個更需要關注？", "如何向一般民眾說明風險？",
    ],
    "場址與停車場": [
        "停車場適合設置嗎？", "地下停車場可以嗎？", "半戶外停車場可以嗎？", "距離車道多遠才安全？", "為什麼要避開倒車熱區？",
        "可以設在出入口旁嗎？", "可以設在逃生門旁嗎？", "可以靠近排水溝嗎？", "可以靠近 EV 充電樁嗎？", "可以靠近建物門窗嗎？",
        "可以設在商場或商辦嗎？", "可以設在社區嗎？", "公有土地可以設置嗎？", "如何處理淹水風險？", "如何處理颱風風險？",
    ],
    "容器與櫃體": [
        "容器如何確認合格？", "如何避免容器翻倒？", "如何保護容器閥件？", "滿桶與空桶可以混放嗎？", "異常桶如何處理？",
        "櫃體需要哪些通風設計？", "櫃體是否需要防爆電氣？", "櫃體是否需要漏氣偵測？", "櫃體是否需要溫度監控？", "櫃體停電怎麼辦？",
        "通訊中斷怎麼辦？", "櫃門卡住怎麼辦？", "可以在櫃內充填 LPG 嗎？", "櫃體需要接地或避雷嗎？", "如何避免日照高溫？",
    ],
    "人員與營運": [
        "誰可以取桶？", "如何避免未成年人或未授權者取桶？", "使用者拿錯桶怎麼辦？", "使用者不會操作怎麼辦？", "補貨人員需要什麼訓練？",
        "補貨時如何避免影響車流？", "多久需要巡檢一次？", "遠端告警由誰處理？", "夜間發生告警怎麼辦？", "如何處理惡意破壞？",
        "如何防止偷竊？", "如何處理客訴？", "如何管理庫存差異？", "如何保存事件紀錄？", "如何進行緊急演練？",
    ],
    "消防、法規與主管機關": [
        "台灣現行法規是否明確允許智取櫃？", "設置前要先找哪個主管機關？", "智慧櫃屬於販賣場所還是儲存場所？", "可以存放多少 LPG？", "是否必須設置漏氣警報？",
        "是否必須設置滅火器？", "是否有禁菸火距離要求？", "為什麼要維持 40°C 以下？", "是否需要申請儲存場所證明？", "容器檢驗期限如何管理？",
        "消防機關若要求改善怎麼辦？", "不同縣市要求會不會不同？", "場址變更是否需要重審？", "設備改版是否需要重新確認？", "國外法規可以直接套用嗎？",
    ],
    "保險、責任與品牌": [
        "需要哪些保險？", "事故責任由誰承擔？", "停車場業者需要承擔什麼？", "如何處理第三人車損或人傷？", "保險公司會關注哪些資料？",
        "為什麼需要退出機制？", "如何避免影響合作夥伴品牌？", "發生社群爭議怎麼辦？", "發生重大事件怎麼通報？", "是否需要公開安全資料？",
    ],
    "商業與試點決策": [
        "為什麼要先做低量試點？", "第一個場址應如何選？", "試點要觀察哪些 KPI？", "什麼情況應立即停用？", "什麼條件下可以擴大？",
        "什麼條件下應退出？", "如何評估單場 Business Case？", "最常見的營運問題是什麼？", "最大的重大事故風險是什麼？", "最大的成功因素是什麼？",
    ],
    "國際最佳實務與台灣化": [
        "國外為什麼敢設置容器交換櫃？", "國外是否普遍使用全自動 LPG 販賣機？", "美國做法的核心是什麼？", "英國做法的核心是什麼？", "德國做法的核心是什麼？",
        "法國做法的核心是什麼？", "澳洲與紐西蘭做法的核心是什麼？", "日本做法的核心是什麼？", "韓國做法的核心是什麼？", "國際最佳實務如何轉成旺來方案？",
    ],
}

def qa_answer(category, q):
    if "地下" in q:
        return ("地下與封閉空間有氣體累積與疏散風險，首波不建議。", "英國、法國與紐西蘭均強調戶外或充分通風；法國明定部分地下儲存禁止。", "列為原則退件；如主管機關認為可評估，仍須專案工程與消防審查。", "UK-01/FR-01/NZ-01")
    if "菸蒂" in q or "爆炸" in q or "火災" in q:
        return ("事故需要洩漏、可燃濃度累積、點火源與未及時控制等條件；不能說絕對不會，也不能把菸蒂視為唯一風險。", "各國均以避免洩漏與累積、控制點火源、通風及緊急應變管理。", "戶外通風、禁菸、漏氣告警、巡檢、異常停用與消防審查。", "US-01/UK-01/TW-02")
    if "EV" in q or "距離" in q or "門窗" in q or "出入口" in q or "逃生" in q or "排水" in q:
        return ("距離不能只用單一數字決定；需看儲量、障壁、通風、點火源、人車動線與主管機關要求。", "英國、法國、紐西蘭均要求避開建物開口、排水、通道、火源或以安全距離／屏障管理。", "以現勘配置圖提交消防與場站審查；未確認前不得施工。", "UK-03/FR-01/NZ-01/TW-02")
    if "保險" in q or "責任" in q or "品牌" in q or "社群" in q:
        return ("風險可控不等於責任消失；須在合約、保險與事件通報中明確分工。", "成熟市場將供應商維護、安全、保全與緊急安排視為營運責任。", "旺來承接設備、容器、補貨、客服與異常；場站保留場域與品牌審查權，並依核保結果配置保險。", "UK-03/JP-01")
    if "法規" in q or "主管機關" in q or "儲存場所" in q or "販賣場所" in q or "存放多少" in q or "申請" in q:
        return ("須依實際儲放量、容器狀態、營運與場址由主管機關判定，不能自行定性。", "各國均依用途、數量及場所採分級管理。", "設置前向轄區消防機關取得書面意見，並建立法規與變更管理 Gate。", "TW-01/TW-03")
    if "容器" in q or "滿桶" in q or "空桶" in q or "異常桶" in q or "閥件" in q:
        return ("容器即使視為空桶仍可能有殘氣；須直立、固定、保護閥件、追溯並隔離異常。", "美國、英國、日本與台灣均重視容器檢驗、直立、防損傷及滿／殘氣管理。", "掃碼綁定檢驗期限；滿桶、殘氣桶與異常桶分流；補貨雙重檢查。", "US-01/UK-02/JP-02/TW-01")
    if "監控" in q or "告警" in q or "停電" in q or "通訊" in q or "櫃門" in q or "漏氣偵測" in q or "溫度" in q:
        return ("智慧功能的價值是提早發現與限制存取，但必須有故障安全與人工應變。", "德國 TRGS 510 強調警報、檢查與營運組織；台灣規範亦重視漏氣警報與溫度。", "設定告警分級、斷線／停電自動停用、遠端值班、現場回應與到達服務時限（SLA）及事件紀錄。", "DE-01/TW-01/TW-02")
    if "補貨" in q or "巡檢" in q or "訓練" in q or "操作" in q or "取桶" in q or "庫存" in q:
        return ("高頻風險多來自操作與搬運，必須以權限、訓練、SOP、巡檢與紀錄降低。", "英國 HSE 指出容器搬運與連接涉及人為錯誤；日本要求供應商保安責任。", "APP 身分驗證、一次一格、補貨認證、雙重檢查、巡檢與異常隔離。", "UK-02/JP-01")
    if "颱風" in q or "淹水" in q or "高溫" in q or "日照" in q:
        return ("極端天候可能造成容器位移、流失、浸水、通訊中斷或溫度上升。", "日本已將淹水區容器防流失納入規範；台灣要求容器保持 40°C 以下。", "淹水圖資篩選、固定、遮陽、溫度告警、颱風前停用與必要撤桶。", "JP-03/TW-02")
    if "國外" in q or "美國" in q or "英國" in q or "德國" in q or "法國" in q or "澳洲" in q or "紐西蘭" in q or "日本" in q or "韓國" in q or "國際" in q:
        return ("國外可運作的前提是把容器交換點視為受控危險物質節點，而不是普通零售設備。", "共通措施包括戶外通風、限制接觸、容器固定、防撞、數量分級、巡檢、供應商責任與主管機關監督。", "將共通原則轉成一場一審、消防預審、低量試點與多層防護。", "US-01/UK-01/DE-01/FR-01/NZ-01/JP-01/KR-01")
    if "試點" in q or "KPI" in q or "Business Case" in q or "擴大" in q or "退出" in q or "停用" in q or "成功" in q or "營運問題" in q:
        return ("先用低量、限時、可退出的試點取得真實風險與營運數據，再決定是否擴大。", "國際管理普遍依數量、場址與風險分級，提高要求。", "以 90 天檢視使用量、告警、撞擊、異常桶、客訴、維運、成本與殘餘風險；每新增一場重新審查。", "NZ-01/DE-01/TW-01")
    return ("應以場址、儲量、設備、操作及責任分工進行個案風險評估，不能只用單一設備特徵判斷。", "國際共通做法是分層控制、持續檢查與可追溯管理。", "納入消防預審、一場一審、監控、SOP、保險及退出機制。", "US-01/UK-02/DE-01/TW-01")

qa_no = 1
for category, questions in qa_topics.items():
    doc.add_heading(category, level=2)
    for q in questions:
        pro, intl, tw, src = qa_answer(category, q)
        doc.add_heading(f"Q{qa_no}. {q}", level=3)
        labeled_paragraph(doc, "A｜", pro, bold=True)
        labeled_paragraph(doc, "國際最佳實務｜", intl)
        labeled_paragraph(doc, "台灣／旺來方案｜", tw, bold=True)
        labeled_paragraph(doc, "主要來源｜", src)
        qa_no += 1

# Source appendix
page_break(doc)
doc.add_heading("九、國際最佳實務 → 旺來智取櫃 → 台灣落地對照表", level=1)
table(doc, ["國際最佳實務", "旺來智取櫃設計／營運", "台灣落地與驗證"], [
    ["戶外與良好通風優先", "櫃體通風、避免密閉與積氣", "場址現勘＋消防確認；地下／低窪首波退件"],
    ["限制公眾與未授權接觸", "APP 身分、一次性權限、鎖控、CCTV", "個資／權限稽核、場站告示與客服"],
    ["防止物理損壞與車撞", "位置退縮、防撞柱、固定基礎", "停車動線圖、工程驗收、場站同意"],
    ["容器直立、固定、閥件保護", "一次一格、定位、掃碼追溯", "容器檢驗有效、補貨與巡檢紀錄"],
    ["依數量提高管理要求", "低量試點、庫存上限、超量鎖定", "消防核定量、變更管理"],
    ["供應商負責檢查與應變", "旺來承接補貨、客服、異常與撤桶", "SOP、服務時限（SLA）、教育訓練與合約責任"],
    ["告警、檢查與紀錄", "門、溫度、漏氣、斷線、CCTV", "告警紀錄、月度稽核、事件改善"],
    ["極端天候與容器流失控制", "固定、遮陽、淹水與颱風停用", "淹水圖資、颱風 SOP、必要撤桶"],
])

doc.add_heading("十、來源索引", level=1)
table(doc, ["編號", "國家／機關", "文件", "網址"], SOURCES, size=7.3)
callout(doc, "來源使用原則", "國際來源用於辨識最佳實務與風險控制，不直接替代台灣法規。法規與官方網站可能更新，正式設置前應重新確認最新版本並取得主管機關意見。")

doc.add_heading("十一、合作夥伴下一步建議", level=1)
manual_numbered(doc, [
    "由旺來提供櫃體、容器、監控、補貨、應變與保險文件包。",
    "合作夥伴提供 1 個戶外、通風、防撞與監控條件良好的候選場址。",
    "雙方共同向轄區消防機關進行概念預審，確認法律定性與要求。",
    "完成 A 級場址評分、工程／保險審查及責任矩陣。",
    "條件成立後啟動低量、限時、可退出的 90 天試點。",
])

doc.save(TARGET)
print(f"target={TARGET}")
print(f"backup={BACKUP}")
print(f"qa_count={qa_no - 1}")
print(f"sources={len(SOURCES)}")
