from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "04_簡版PPT" / "VIVI_商開會議_報告口稿與應答手冊.docx"
ORANGE = "F05A24"
LIGHT_ORANGE = "FFF1EA"
DARK = "202020"
GRAY = "666666"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        shade(table.rows[0].cells[i], ORANGE)
        set_cell_text(table.rows[0].cells[i], header, bold=True, color="FFFFFF", size=9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9)
            if len(table.rows) % 2 == 0:
                shade(cells[i], "FFF9F6")
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, LIGHT_ORANGE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(ORANGE)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_script(doc, slide, title, purpose, script, transition, caution=None):
    doc.add_heading(f"第 {slide} 頁｜{title}", level=2)
    add_callout(doc, "這頁目的", purpose)
    p = doc.add_paragraph()
    r = p.add_run("建議口稿\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(ORANGE)
    p.add_run(script)
    p.paragraph_format.space_after = Pt(6)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("轉場句：")
    r2.bold = True
    p2.add_run(transition)
    if caution:
        p3 = doc.add_paragraph()
        r3 = p3.add_run("注意：")
        r3.bold = True
        r3.font.color.rgb = RGBColor.from_string("C00000")
        p3.add_run(caution)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.7)
section.bottom_margin = Cm(1.7)
section.left_margin = Cm(1.9)
section.right_margin = Cm(1.9)

styles = doc.styles
styles["Normal"].font.name = "Microsoft JhengHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].paragraph_format.space_after = Pt(5)
for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
    styles[style_name].font.name = "Microsoft JhengHei"
    styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
styles["Heading 1"].font.color.rgb = RGBColor.from_string(ORANGE)
styles["Heading 2"].font.color.rgb = RGBColor.from_string(DARK)
styles["Heading 3"].font.color.rgb = RGBColor.from_string(ORANGE)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("ViVi PARK × 旺來\n商開會議報告口稿與應答手冊")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor.from_string(ORANGE)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("會前複習、逐頁口稿、合作方案、商業模型與現場 Q&A\n會議日期：2026 年 6 月 11 日 15:00").italic = True
doc.add_paragraph()
add_callout(
    doc,
    "這份文件怎麼用",
    "會議前先讀「一頁速記」與「不能直接答應的事項」；報告前複習逐頁口稿；現場遇到商業、分潤、智慧櫃或風險問題時，直接查閱對應章節。",
)

doc.add_heading("一頁速記：今天要談成什麼", level=1)
add_table(
    doc,
    ["項目", "會議目標"],
    [
        ["核心定位", "旺來成為 ViVi 生活服務生態系中的家庭能源品類合作夥伴。"],
        ["第一階段", "先談 APP 上架、票券／序號核銷、會員導流與成交分潤。"],
        ["第二階段", "智慧櫃體採一場一審，由營運／風險單位另行評估。"],
        ["希望取得", "確認首波品項、上架規格、核銷流程、分潤計算、PoC KPI 與下一步窗口。"],
        ["不要變成", "不要只談放櫃體、不要當場無條件接受 20%、不要把推估會員數當官方事實。"],
    ],
    [3.0, 13.5],
)

doc.add_heading("會議開場與收尾", level=1)
doc.add_heading("建議開場", level=2)
doc.add_paragraph(
    "Phoebe 您好，也謝謝 ViVi PARK 商開團隊安排這次會議。今天我們希望先承接貴司回信提到的會員導流、APP 商品或服務上架、票券核銷與分潤合作。旺來希望討論的不是單純放置設備，而是家庭能源服務能否成為 ViVi 生活服務生態系中，一個可商品化、可核銷、可履約，也可透過數據驗證的新服務品類。"
)
doc.add_heading("建議收尾", level=2)
doc.add_paragraph(
    "今天我們希望先確認一個低風險、可量測、可回顧的第一階段合作模型。若方向可行，旺來會依 ViVi 的上架規格補齊首波商品與服務清單、核銷方式、素材需求及 90 天 PoC 指標；智慧櫃體則依一場一審原則，另提供風險與營運資料供相關單位評估。"
)

doc.add_heading("逐頁報告口稿", level=1)
add_script(
    doc, 1, "ViVi 可把停車會員，升級為生活服務交易入口",
    "先把合作定位說清楚，避免對方以為旺來只想進場放設備。",
    "ViVi 已經具備 APP、會員、Vi 幣、票券與線下場站。旺來可以補上的，是家庭能源這個具明確需求、可回購、也能由旺來承接履約的生活服務品類。今天我們想先討論：旺來服務能否透過 ViVi APP 成為新的交易與會員互動場景；智慧櫃則作為另一條可獨立評估的線下合作線。",
    "因此，我們先從 ViVi 已經具備的平台資產開始看。",
)
add_script(
    doc, 2, "ViVi 真正的資產不是車位，而是會員、支付、票券與場站觸點",
    "表達旺來理解 ViVi 的平台價值，而不是把對方當場地主。",
    "我們對 ViVi 的理解是：核心資產不只在停車格，而在會員關係、Vi 幣與票券交易能力、停車與生活服務流量入口，以及線下場站觸點。這使 ViVi 有條件把新的生活服務品類快速帶到車主會員面前，再以實際交易判斷是否值得放大。",
    "既然 ViVi 是平台，下一個問題就是：旺來可以補上什麼品類缺口？",
    "300+、700 場站等數字以 ViVi 官方最新口徑為準；不要主動宣稱未確認會員總數。",
)
add_script(
    doc, 3, "家庭能源能補上 ViVi 生態圈中的剛性、可回購品類缺口",
    "說明家庭能源不是一般優惠券，而是可能形成回購與履約關係的品類。",
    "ViVi 已經透過停車、點數與票券創造會員互動。家庭能源的不同，在於需求明確、具回購性，且可以形成客服、配送、自取與長期服務關係。若第一階段的票券與服務核銷成立，ViVi 的平台飛輪就能增加一個非停車型的生活服務交易場景。",
    "這個新場景對 ViVi 的價值，可以拆成五個具體方向。",
)
add_script(
    doc, 4, "Revenue／Traffic／Retention／Differentiation／ESG",
    "用五個價值槓桿回答「ViVi 為什麼要合作」。",
    "我們把合作價值拆成五個面向。Revenue 是成交後的分潤；Traffic 是雙方會員互相接觸；Retention 是增加會員回訪與持續使用理由；Differentiation 是讓 ViVi 從智慧停車延伸為車主生活服務平台；ESG 則是數位票券、安全治理與未來可管理線下節點的案例。五項不是五個專案，而是衡量合作值不值得繼續的五種價值。",
    "合作可能很多，但第一階段必須收斂。",
    "ESG 不要講成主要收益；它是加分與治理價值。",
)
add_script(
    doc, 5, "30+ 合作想法收斂成三條第一階段主線",
    "展現有完整思考，但不讓會議失焦。",
    "會員、支付、點數、票券、通路、場站、智慧櫃、數據、ESG 與 B2B 都有合作可能。但若第一階段全部展開，執行成本會太高。因此建議先收斂為會員導流、票券／點數與 APP 特約商店上架三條主線；智慧櫃則採一場一審，由另一條營運與風險流程評估。",
    "接著用影響與執行難度，決定先做什麼。",
)
add_script(
    doc, 6, "先選高影響、低投入的驗證項目",
    "讓合作順序有邏輯，而不是空泛列出券種。",
    "第一波建議先測試首購券，確認 ViVi 會員是否願意成為旺來新客；第二步用會員專屬券，確認 ViVi 會員權益是否能提高領券與核銷；第三步才測補充服務券，確認是否形成回購。智慧櫃可同步提供風險資料與候選場站條件，但實際設置採一場一審。",
    "如果先從票券開始，交易與核銷流程需要怎麼設計？",
)
add_script(
    doc, 7, "票券／序號模式與 20% 分潤",
    "把雙方責任、交易流程與分潤討論講清楚。",
    "建議流程是 ViVi APP 上架商品或服務券，會員購買或領取後取得序號，再到旺來平台完成資格確認、核銷與履約，月底依有效成交資料對帳。ViVi 提出的 20% 可作為討論基礎，但需要先確認計算基礎、折扣、稅金、退款、配送費及有效成交定義，並依不同品項毛利評估。",
    "交易模型成立後，才有條件精準設計會員導流。",
    "不要當場承諾所有品項固定 20%；可以說『以 20% 作為討論基礎，需依品項與履約成本確認』。",
)
add_script(
    doc, 8, "會員導流先用可驗證假設",
    "避免用沒有根據的會員總數或人口統計說服對方。",
    "目前已知 ViVi 會員集中雙北與高雄、旺來會員集中台中，這是雙方區域互補的基礎。但第一波不應假設全體會員都會轉換，而應請 ViVi 協助定義可測試會員池，例如近 90 天活躍會員、特定區域會員或曾使用生活服務的會員，再觀察曝光、點擊、領券與成交。",
    "因此第九頁的數字不是營收承諾，而是 PoC 的測量方法。",
)
add_script(
    doc, 9, "90 天 PoC 敏感度模型",
    "用簡單數學說明要收集什麼數據，以及如何判斷合作是否值得繼續。",
    "第九頁以每 10,000 名可觸達會員為測算單位。先看有多少人點擊，再看點擊者中有多少人成交。成交訂單乘以平均客單價，就是 GMV；GMV 再乘以雙方約定的分潤比例，才是 ViVi 可能取得的分潤。這些是敏感度情境，不是預測；目的是在 PoC 前先把數據欄位與決策標準講清楚。",
    "除了線上交易，另一條合作線是智慧櫃的一場一審。",
)
add_script(
    doc, 10, "智慧櫃體進入一場一審的 Business Case",
    "回應對方對電力、菸蒂、安全與場站營運的疑問，並取得下一步窗口。",
    "智慧櫃體不會以所有場站通用的方式推進，而是每一個候選場站分別確認通風、供電、動線、防撞、監控、保險、權責與現場管理。旺來負責提供設備、安全、補貨、客服、保險與異常處理資料；ViVi 保留場站管理與品牌審查權。今天希望確認是否能由商開協助轉交風險資料，並安排營運或風險窗口進行下一步評估。",
    "最後請雙方確認首波品項、交易流程、分潤規則、PoC 與智慧櫃評估窗口。",
)

doc.add_heading("商業名詞白話解釋", level=1)
add_table(
    doc,
    ["名詞", "白話意思", "本案怎麼用"],
    [
        ["GMV", "透過 ViVi 成功成交的訂單總金額，不等於旺來營收或獲利。", "成交訂單數 × 平均客單價。"],
        ["20% 分潤", "每筆有效成交中，ViVi 希望取得成交金額的 20%。", "需確認計算基礎、退款、折扣、配送費、稅金與品項差異。"],
        ["CPA", "按指定成果付費，例如每取得一名有效新會員支付固定費用。", "適合純會員導流，但需定義什麼是有效會員。"],
        ["CPS", "實際成交後才按比例分潤。", "ViVi 信中的 20% 分潤較接近 CPS。"],
        ["CTR", "看到活動的人之中，有多少人點擊。", "點擊人數 ÷ 曝光人數。"],
        ["轉換率", "點擊或領券的人之中，有多少人最後成交。", "成交訂單 ÷ 點擊或領券人數。"],
        ["AOV", "平均每筆訂單成交多少錢。", "GMV ÷ 成交訂單數。"],
        ["PoC", "小規模驗證，不是一開始全面合作。", "用 90 天測試交易、核銷、會員反應與履約成本。"],
        ["Business Case", "針對單一場站，判斷收益、成本、風險與責任是否合理。", "智慧櫃採一場一審。"],
    ],
    [2.4, 7.0, 7.0],
)

doc.add_heading("20% 分潤：現場怎麼談", level=1)
doc.add_paragraph("ViVi 信中的概念可理解為：ViVi 提供 APP、會員流量、上架與交易入口；旺來提供商品、核銷、履約、客服與售後；有效成交後，ViVi 希望取得 20%。")
add_table(
    doc,
    ["範例", "金額"],
    [
        ["一筆服務券成交價", "1,000 元"],
        ["ViVi 若按成交價取得 20%", "200 元"],
        ["旺來帳面取得", "800 元"],
        ["旺來仍需負擔", "商品、配送、客服、退款、營運與稅務成本"],
    ],
    [8.0, 8.0],
)
add_callout(
    doc,
    "建議回答",
    "我們理解 ViVi 以 20% 作為特約商店合作期待。旺來願意以此作為討論基礎，但不同品項的商品成本、配送與履約責任不同，希望先確認計算基礎與有效成交定義，再共同選出適合測試的品項。",
)
doc.add_heading("20% 分潤必問八題", level=2)
for item in [
    "20% 以原價、折扣後售價、含稅或未稅金額計算？",
    "購買即算成交，還是完成核銷／履約才算成交？",
    "優惠折扣由 ViVi、旺來或雙方共同負擔？",
    "金流手續費由誰負擔？",
    "退款、取消、未核銷票券如何處理？",
    "配送費是否也納入 20% 分潤？",
    "實體商品、服務券與會員導流是否適用相同分潤？",
    "是否可依品項毛利設定不同分潤或採階梯式分潤？",
]:
    add_number(doc, item)

doc.add_heading("五大價值槓桿的本質", level=1)
add_table(
    doc,
    ["槓桿", "本質問題", "ViVi 可能得到", "建議 KPI"],
    [
        ["Revenue", "合作能不能增加收入？", "商品／服務成交分潤。", "GMV、成交數、分潤收入"],
        ["Traffic", "合作能不能帶來更多可接觸客戶？", "雙北／高雄與台中會員互導。", "曝光、點擊、新會員"],
        ["Retention", "會員會不會更常回來？", "增加 APP、票券與生活服務使用理由。", "回訪率、回購率、核銷率"],
        ["Differentiation", "ViVi 是否變得更難被取代？", "從智慧停車延伸為車主生活服務平台。", "新品類數、會員使用場景"],
        ["ESG", "是否形成治理、安全或永續案例？", "數位票券、安全可管理的能源服務節點。", "紙本減量、異常率、治理文件"],
    ],
    [2.5, 4.3, 6.0, 3.5],
)

doc.add_heading("四階段合作執行方案", level=1)
add_table(
    doc,
    ["階段", "目的", "具體方案", "執行工作", "判斷指標"],
    [
        ["1. 首購券", "測試 ViVi 會員是否願意第一次使用旺來。", "新戶首購優惠、首次服務諮詢、指定區域新客方案。", "ViVi 曝光／發券；旺來核銷、資格確認、履約。", "點擊、領券、首購、獲客成本"],
        ["2. 會員專屬券", "增加 ViVi 會員權益與平台價值。", "ViVi 限定價格、限定贈品、Vi 幣折抵或回饋。", "共同設計會員限定素材與活動規則。", "領券率、核銷率、Vi 幣使用、APP 回訪"],
        ["3. 補充服務券", "驗證是否形成回購與持續交易。", "補充服務優惠、回購券、預約配送／自取方案。", "識別既有客戶、建立回購與客服流程。", "回購率、回購間隔、會員終身價值、履約成本"],
        ["4. 智慧櫃", "建立線下自取、展示與履約節點。", "依場站條件設置智慧櫃或展示節點。", "每場確認通風、供電、動線、防撞、監控、保險與權責。", "使用量、異常率、補貨效率、客訴、場站影響"],
    ],
    [2.1, 3.5, 4.7, 4.7, 3.8],
)

doc.add_heading("建議交易流程與責任分工", level=1)
for step in [
    "ViVi APP 上架旺來商品、服務券或會員優惠。",
    "ViVi 會員購買、領券或取得折扣／服務序號。",
    "會員前往旺來指定頁面，輸入序號並確認服務資格、地區與需求。",
    "旺來完成核銷、客服、配送、自取或其他履約。",
    "旺來回傳有效成交、退款與核銷資料。",
    "雙方每月依約定口徑對帳，ViVi 取得分潤。",
]:
    add_number(doc, step)
add_table(
    doc,
    ["ViVi 主要負責", "旺來主要負責", "雙方共同確認"],
    [
        ["APP 上架與會員曝光", "商品／服務設計與供應", "上架規格與審稿"],
        ["票券、訂單或序號資訊", "序號核銷與資格確認", "有效成交定義"],
        ["行銷素材投放與會員入口", "客服、配送、自取、售後與退款處理", "分潤、對帳與資料欄位"],
        ["月結資料與平台規則", "回傳成交、核銷及異常資料", "個資、品牌與客訴處理"],
    ],
    [5.3, 5.3, 5.3],
)

doc.add_heading("第九頁完整解讀：數字怎麼算", level=1)
add_callout(
    doc,
    "核心公式",
    "可觸達會員數 × 點擊率 × 成交轉換率 = 成交訂單數；成交訂單數 × 平均客單價 = GMV；GMV × 分潤比例 = ViVi 分潤。",
)
add_table(
    doc,
    ["基本情境範例", "計算"],
    [
        ["可觸達會員", "10,000 人"],
        ["點擊率 4%", "400 人點擊"],
        ["點擊後成交率 10%", "40 筆成交"],
        ["平均客單價 1,000 元", "GMV = 40,000 元"],
        ["若 ViVi 分潤 20%", "ViVi 分潤 = 8,000 元"],
        ["旺來帳面取得", "32,000 元，仍需扣除商品與履約成本"],
    ],
    [8.0, 8.0],
)
doc.add_paragraph(
    "第九頁不是承諾一定成交 40 筆，而是告訴雙方 PoC 必須蒐集哪些資料。真正要比較的是：不同素材、品項、區域與會員分群，哪一種能以合理成本帶來成交與回購。"
)

doc.add_heading("現場可能被問的問題與建議回答", level=1)
qa_rows = [
    ["為什麼 ViVi 會員會需要旺來？", "家庭能源具有明確需求與回購特性；但不預設所有會員都需要，建議以區域與活躍會員池先做小規模驗證。"],
    ["為什麼不直接放智慧櫃？", "智慧櫃可以獨立一場一審；先確認場站條件、會員需求與風險控制，避免設備設置後使用不足。"],
    ["20% 分潤能不能接受？", "願意以 20% 作為討論基礎，但需依品項毛利、配送與履約成本確認計算基礎與有效成交定義。"],
    ["旺來能上架什麼？", "可先從新客首購、會員專屬優惠與服務券中選擇 1 至 3 項，再依 ViVi 上架規格調整。"],
    ["核銷與客服誰處理？", "旺來可承接序號核銷、資格確認、客服、配送、自取、售後與異常處理。"],
    ["ViVi 會員大約多少？", "公開資料未揭露官方總數；為設計 PoC，更重要的是可觸達與近 90 天活躍會員池，建議由 ViVi 提供可測試口徑。"],
    ["菸蒂會不會造成爆炸？", "不能用一句『不會』概括。風險取決於是否有洩漏、濃度、通風與點火源；因此需以設備防護、通風、偵測、禁菸、防撞及場站審查控制。"],
    ["供電怎麼處理？", "每個場站先確認合法接電方式、電箱距離、用電負載、費用與權責；條件不成立就不設置。"],
    ["發生事故或客訴誰負責？", "設備、商品、補貨、客服、售後與異常由旺來承接；場站管理與品牌審查由 ViVi 保留，並事前約定事件分級、通報與保險。"],
    ["合作成功後怎麼放大？", "先看 90 天曝光、核銷、成交、回購與客服成本；數據成立後再增加品項、會員分群或評估更多場站。"],
]
add_table(doc, ["可能問題", "建議回答"], qa_rows, [5.0, 11.2])

doc.add_heading("智慧櫃一場一審：必備評估框架", level=1)
add_table(
    doc,
    ["評估項目", "需確認內容", "原則"],
    [
        ["通風與環境", "是否為適合位置、通風條件、是否存在積聚風險。", "條件不明確不設置。"],
        ["供電", "合法接電、電箱距離、負載、費用與權責。", "每場獨立確認。"],
        ["動線與防撞", "不占車格、不擋車道、不影響行人，具必要防撞。", "現勘後決定。"],
        ["監控與異常", "監控、偵測、斷電、警示、客服與通報流程。", "事前定義 SOP。"],
        ["保險與責任", "產品責任、公共意外、設備及營運責任。", "合作前確認文件。"],
        ["品牌與客訴", "現場告示、對外說法、客訴與危機處理。", "ViVi 保留品牌審查權。"],
    ],
    [3.5, 8.0, 4.7],
)

doc.add_heading("會議中必須問清楚的問題", level=1)
for item in [
    "ViVi APP／特約商店目前可上架哪些商品與服務形式？",
    "票券、折扣碼、服務序號與實體商品轉單，各自的標準流程是什麼？",
    "20% 分潤的計算基礎、退款、折扣、稅金、配送費與月結規則為何？",
    "第一波適合測試哪一個會員池：全體、近 90 天活躍、特定區域或特定行為會員？",
    "ViVi 能提供哪些曝光位置、推播或會員溝通工具？",
    "PoC 可取得哪些數據：曝光、點擊、領券、核銷、成交、退款與會員分群？",
    "首波建議上架 1 至 3 個品項，ViVi 認為哪類最容易通過？",
    "智慧櫃風險資料應由哪一個營運／風險窗口承接？",
    "是否可先選 1 個候選場站進行文件與現勘初審？",
    "會後雙方下一步、負責人與預計時間點為何？",
]:
    add_number(doc, item)

doc.add_heading("不能直接答應或說死的事項", level=1)
for item in [
    "不要直接答應所有品項固定 20% 分潤。",
    "不要把 GMV 說成旺來獲利。",
    "不要宣稱 ViVi 官方會員數；公開資料沒有正式揭露。",
    "不要承諾所有停車場都適合設櫃。",
    "不要以『菸蒂絕對不會造成問題』作為風險回答；應說明風險條件與控制措施。",
    "不要承諾未經確認的 API、金流、個資交換或 Vi 幣整合能力。",
    "不要讓智慧櫃議題取代 APP 上架、會員導流與交易合作討論。",
]:
    add_bullet(doc, item)

doc.add_heading("會議後應立即整理的紀錄", level=1)
add_table(
    doc,
    ["紀錄項目", "會後需填寫"],
    [
        ["ViVi 參與者與窗口", "姓名、職稱、負責範圍、聯絡方式"],
        ["首波品項", "可上架／需修改／不可上架"],
        ["交易與核銷", "平台流程、序號規格、資料欄位"],
        ["分潤與對帳", "比例、基礎、退款、月結、稅務"],
        ["會員導流", "可用會員池、曝光工具、素材規格"],
        ["PoC", "期間、KPI、資料提供方式、成功門檻"],
        ["智慧櫃", "風險窗口、候選場站、需補文件"],
        ["下一步", "雙方負責人、交付內容、日期"],
    ],
    [5.0, 11.2],
)

doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = end.add_run("誠懇・專業・當責・共榮")
run.bold = True
run.font.color.rgb = RGBColor.from_string(ORANGE)
run.font.size = Pt(12)

doc.save(OUTPUT)
print(OUTPUT)
